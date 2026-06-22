from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)

# ── Palette ───────────────────────────────────────────────────────────────────
COLOR_NAVY   = RGBColor(0x1B, 0x2A, 0x4A)   # headings
COLOR_SLATE  = RGBColor(0x3A, 0x4A, 0x6B)   # sub-headings
COLOR_TEAL   = RGBColor(0x00, 0x7A, 0x87)   # accent / labels
COLOR_BODY   = RGBColor(0x1A, 0x1A, 0x1A)   # body text
COLOR_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_HEADER = RGBColor(0x1B, 0x2A, 0x4A)   # table header bg
COLOR_ROW1   = RGBColor(0xF2, 0xF5, 0xFA)   # alternating row tint
COLOR_BORDER = RGBColor(0xC5, 0xCE, 0xDE)
COLOR_AMBER  = RGBColor(0xB8, 0x6E, 0x00)

# label badge colours
BADGE = {
    "no_hallucination":       RGBColor(0x15, 0x7A, 0x3E),
    "minor_hallucination":    RGBColor(0xB8, 0x6E, 0x00),
    "major_hallucination":    RGBColor(0xA8, 0x1C, 0x1C),
    "unknown_not_computable": RGBColor(0x4A, 0x4A, 0x6B),
}

LABEL_TEXT = {
    "no_hallucination":       "No Hallucination",
    "minor_hallucination":    "Minor Hallucination",
    "major_hallucination":    "Major Hallucination",
    "unknown_not_computable": "Unknown – Not Computable",
}

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = str(rgb)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="C5CEDE", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_paragraph_border_bottom(para, color="C5CEDE"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def add_paragraph_border_left(para, color="007A87", sz="12"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    sz)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    add_paragraph_border_bottom(p, "1B2A4A")
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_NAVY
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_SLATE
    return p

def req_id_para(req_id, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(2)
    add_paragraph_border_bottom(p, "007A87")
    r1 = p.add_run(req_id + "  ")
    r1.bold      = True
    r1.font.size = Pt(15)
    r1.font.color.rgb = COLOR_TEAL
    r2 = p.add_run(title)
    r2.bold      = True
    r2.font.size = Pt(15)
    r2.font.color.rgb = COLOR_NAVY
    return p

def body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_BODY
    return p

def field_label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_TEAL
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after   = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_BODY
    return p

def add_examples_table(rows):
    """rows = list of (label_key, example_text)"""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Outcome"
    hdr[1].text = "Example"
    for i, cell in enumerate(hdr):
        set_cell_bg(cell, COLOR_HEADER)
        set_cell_borders(cell, "1B2A4A", "6")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_WHITE

    for idx, (label_key, example_text) in enumerate(rows):
        row   = table.add_row()
        cells = row.cells
        bg    = COLOR_ROW1 if idx % 2 == 0 else COLOR_WHITE

        set_cell_bg(cells[0], bg)
        set_cell_borders(cells[0])
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p0 = cells[0].paragraphs[0]
        run0 = p0.add_run(LABEL_TEXT[label_key])
        run0.bold = True
        run0.font.size = Pt(9.5)
        run0.font.color.rgb = BADGE[label_key]

        set_cell_bg(cells[1], bg)
        set_cell_borders(cells[1])
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        run1 = p1.add_run(example_text)
        run1.font.size = Pt(9.5)
        run1.font.color.rgb = COLOR_BODY

    for row in table.rows:
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(4.8)

    doc.add_paragraph()

def callout_body(text):
    """Body text with a teal left border — for notes and callouts."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.2)
    add_paragraph_border_left(p, "007A87", "12")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_BODY
    return p


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Hallucination Evaluation Requirements")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = COLOR_NAVY

p2 = doc.add_paragraph()
r2 = p2.add_run("Agentic Observation Platform — Multi-Industry Specification")
r2.font.size = Pt(13)
r2.font.color.rgb = COLOR_SLATE

p3 = doc.add_paragraph()
r3 = p3.add_run("Okahu  ·  Version 1.0  ·  April 2026")
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x7A, 0x8A, 0xA0)
add_paragraph_border_bottom(p3, "C5CEDE")

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — OVERVIEW AND PURPOSE
# ══════════════════════════════════════════════════════════════════════════════

heading1("Overview and Purpose")

body(
    "This document defines the functional requirements for a hallucination evaluation module "
    "within an agentic observation platform. Requirements are designed to apply across any "
    "agentic application regardless of industry, agent framework, or topology — including "
    "financial services, healthcare, legal, logistics, retail, and other regulated and "
    "non-regulated domains."
)

body(
    "Requirements are grounded in trace analysis of a reference implementation: the LG Travel "
    "Agent, a LangGraph multi-agent application using a supervisor pattern with sub-agents for "
    "flight booking, hotel booking, and weather retrieval, instrumented via Monocle/OpenTelemetry. "
    "All six production traces analysed were labelled no_hallucination by the prior evaluation "
    "template. Hallucinations were present in all six. The reference implementation is documented "
    "in the Appendix."
)

body(
    "Where requirements depend on platform capabilities not yet available in the current release, "
    "those dependencies are explicitly identified and captured in the Platform Roadmap section "
    "as Phase 2 enhancements."
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — OUTCOME LABEL DEFINITIONS
# ══════════════════════════════════════════════════════════════════════════════

heading1("Outcome Label Definitions")
heading2("Label Descriptions")

label_defs = [
    (
        "no_hallucination",
        "The agent's response is fully supported by tool outputs and input context. No fabricated "
        "or unsupported claims are present. Standardised, lossless transformations — such as "
        "capitalisation normalisation, accepted abbreviations, and well-known code expansions — "
        "are included in this category. These have only one correct answer and introduce no new information."
    ),
    (
        "minor_hallucination",
        "The agent's response contains a small but real factual deviation — information that has a "
        "truthful or untruthful answer and was not present in the source. This requires actual content "
        "the agent introduced, not formatting or standard abbreviation. Examples: an inferred unit not "
        "present in the tool output, a dropped accent on a proper noun, a qualifying detail omitted in "
        "a relay, or a value computed from source data that could be incorrect. The deviation does not "
        "materially mislead the user but is a verifiable departure from the source."
    ),
    (
        "major_hallucination",
        "The agent's response contains a substantive fabrication, substitution, or contradiction that "
        "could cause the user to take incorrect action. Includes: tool outputs overridden with different "
        "facts, specific entities invented with no source, empty tool results treated as successful "
        "confirmations, conclusions contradicted across agents in the same session, and definitive "
        "claims made without evidentiary support."
    ),
    (
        "unknown_not_computable",
        None  # special — three sub-bullets follow
    ),
]

for key, definition in label_defs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    rb = p.add_run(f"{LABEL_TEXT[key]}:  ")
    rb.bold = True
    rb.font.size = Pt(10.5)
    rb.font.color.rgb = BADGE[key]
    if definition:
        rd = p.add_run(definition)
        rd.font.size = Pt(10.5)
        rd.font.color.rgb = COLOR_BODY
    else:
        rd = p.add_run(
            "A determination cannot be made. This label has three distinct sources — "
            "each is actionable in a different way:"
        )
        rd.font.size = Pt(10.5)
        rd.font.color.rgb = COLOR_BODY

# unknown sub-bullets
bullet(
    "Template parameter gap — The evaluation template has no field to assess this check. "
    "Every unknown of this type is a direct signal that a new evaluation parameter is required."
)
bullet(
    "Registry prerequisite not configured — The check requires information that cannot be "
    "derived from span data: a tool response contract (REQ-04, requires Tool Contract Registry) "
    "or an authorized scope declaration (REQ-08, requires Authorized Scope Registry). "
    "See Phase 2 Platform Roadmap for these enhancements. This resolves when the operator "
    "configures the relevant registry entry."
)
bullet(
    "Insufficient trace data — The specific span data needed to make a determination is absent "
    "or unresolvable in the trace."
)

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — PLATFORM CONFIGURATION REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════

heading1("Platform Configuration Requirements")

body(
    "The following information must be available to the evaluation system per agent deployment. "
    "Most is derivable from fully captured trace data. Items marked Registry Required cannot be "
    "derived from spans and require explicit operator configuration."
)

config_rows = [
    (
        "Instrumentation completeness — all agent, tool, and inference spans captured",
        "Trace data (automatic when Monocle/OTel instrumentation is deployed)",
        "All REQs",
        "If spans are missing, affected checks return unknown (insufficient trace data)",
        False,
    ),
    (
        "Tool registry — list of tools available to each agent",
        "Derived from agentic.tool.invocation spans",
        "REQ-01, REQ-02",
        "Derivable from traces when instrumentation is complete",
        False,
    ),
    (
        "Entity schema — named entity types relevant to the deployment domain",
        "REGISTRY REQUIRED — not currently a platform capability",
        "REQ-03 Phase 2, REQ-10 Phase 2",
        "Phase 2 platform enhancement. See Platform Roadmap.",
        True,
    ),
    (
        "Action registry — consequential actions the agent may claim to complete",
        "Derived from turn_end spans and tool invocation spans",
        "REQ-01",
        "Derivable from traces when instrumentation is complete",
        False,
    ),
    (
        "Agent topology — graph of agent-to-agent handoff boundaries",
        "Derived from parent_id, entity.from_agent, scope.agentic fields in spans",
        "REQ-06, REQ-07",
        "Reconstructable from fully captured traces without separate discovery",
        False,
    ),
    (
        "Tool Contract Registry — expected response schema and valid empty states per tool",
        "REGISTRY REQUIRED — not derivable from spans",
        "REQ-04 Phase 2",
        "Phase 2 platform enhancement. See Platform Roadmap.",
        True,
    ),
    (
        "Authorized Scope Registry — permitted data access and action scope per agent",
        "REGISTRY REQUIRED — not derivable from spans",
        "REQ-08 Phase 2",
        "Phase 2 platform enhancement. See Platform Roadmap.",
        True,
    ),
]

config_table = doc.add_table(rows=1, cols=4)
config_table.alignment = WD_TABLE_ALIGNMENT.LEFT
config_table.style = "Table Grid"

config_headers = ["Configuration Item", "Source", "Required For", "Notes"]
hdr = config_table.rows[0].cells
for i, h in enumerate(config_headers):
    set_cell_bg(hdr[i], COLOR_HEADER)
    set_cell_borders(hdr[i], "1B2A4A", "6")
    p = hdr[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = COLOR_WHITE

for idx, (item, source, req_for, notes, registry_req) in enumerate(config_rows):
    row = config_table.add_row()
    cells = row.cells
    bg = COLOR_ROW1 if idx % 2 == 0 else COLOR_WHITE

    for ci, text in enumerate([item, source, req_for, notes]):
        set_cell_bg(cells[ci], bg)
        set_cell_borders(cells[ci])
        cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cells[ci].paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_BODY
        # Source column — highlight registry required rows
        if ci == 1 and registry_req:
            run.bold = True
            run.font.color.rgb = COLOR_AMBER

for row in config_table.rows:
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(2.0)
    row.cells[2].width = Inches(1.1)
    row.cells[3].width = Inches(1.3)

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — EVALUATION TEMPLATE GAP ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

heading1("Evaluation Template Gap Analysis")

body(
    "The current hallucination evaluation template (Evaluation.json) provides a partial "
    "foundation. The table below maps existing parameters against the requirements defined "
    "in this document, identifying what is currently covered, what is partially covered, "
    "and what is missing entirely."
)

gap_rows = [
    (
        "label (enums: no_hallucination, minor_hallucination, major_hallucination)",
        "Basic classification",
        "Missing unknown_not_computable enum; no phase distinction",
        "All REQs",
    ),
    (
        "explanation",
        "Narrative description",
        "Generic — no structured field per check type",
        "All REQs",
    ),
    (
        "hallucination_score",
        "Numeric severity",
        "Scalar — cannot represent per-requirement scores or phase",
        "All REQs",
    ),
    (
        "factual_alignments",
        "Claims supported by context documents",
        "Checks against context documents only — not against tool span outputs",
        "REQ-05 partial",
    ),
    (
        "contradictions",
        "Unsupported or contradicted claims",
        "Single response only — no cross-agent or cross-turn comparison",
        "REQ-05, REQ-06 partial",
    ),
    (
        "hallucination_types (factual_inaccuracy, unsupported_claim, contradiction, fabrication, exaggeration)",
        "Type classification",
        "No types for: tool output faithfulness, action fabrication, scope drift, entity accuracy, "
        "uncertainty from empty result",
        "REQ-05 partial",
    ),
    (
        "context_coverage",
        "Response stays within context bounds",
        "No comparison of tool input parameters against user request specificity",
        "REQ-08 partial",
    ),
    (
        "factual_accuracy",
        "Accuracy against context",
        "Checks context documents — not tool output spans",
        "REQ-05 partial",
    ),
    (
        "verification_status",
        "Claim verification status",
        "Claim-level only — no span-level entity sourcing",
        "REQ-10 partial",
    ),
    (
        "confidence_level",
        "Confidence score",
        "Measures evaluator confidence — not agent confidence calibration relative to evidence",
        "REQ-09 wrong direction",
    ),
    (
        "(missing entirely)",
        "—",
        "No parameter to map claimed actions to tool invocation spans",
        "REQ-01",
    ),
    (
        "(missing entirely)",
        "—",
        "No parameter to compare tool output to agent response; no multi-turn tracking; "
        "no entity-level sourcing; no tool payload inspection for uncertainty",
        "REQ-02, REQ-03, REQ-04, REQ-07",
    ),
]

gap_table = doc.add_table(rows=1, cols=4)
gap_table.alignment = WD_TABLE_ALIGNMENT.LEFT
gap_table.style = "Table Grid"

gap_headers = ["Current Parameter", "Partial Coverage", "Gap", "REQ Mapping"]
hdr = gap_table.rows[0].cells
for i, h in enumerate(gap_headers):
    set_cell_bg(hdr[i], COLOR_HEADER)
    set_cell_borders(hdr[i], "1B2A4A", "6")
    p = hdr[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = COLOR_WHITE

for idx, (param, coverage, gap, req_map) in enumerate(gap_rows):
    row = gap_table.add_row()
    cells = row.cells
    bg = COLOR_ROW1 if idx % 2 == 0 else COLOR_WHITE

    missing = param.startswith("(missing")

    for ci, text in enumerate([param, coverage, gap, req_map]):
        set_cell_bg(cells[ci], bg)
        set_cell_borders(cells[ci])
        cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cells[ci].paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_BODY
        if missing and ci == 0:
            run.italic = True
            run.font.color.rgb = RGBColor(0xA8, 0x1C, 0x1C)

for row in gap_table.rows:
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(1.4)
    row.cells[2].width = Inches(2.2)
    row.cells[3].width = Inches(1.0)

doc.add_paragraph()

body(
    "Parameters required to be added to the evaluation template to implement the requirements "
    "in this document: tool_invocation_verified, tool_output_match, entity_accuracy_check, "
    "tool_payload_inspection, agent_confidence_vs_evidence, cross_agent_consistency, "
    "cross_turn_consistency, scope_drift_detection, source_span_mapping, and "
    "unknown_not_computable as an enum value with source classification "
    "(template_gap, registry_not_configured, insufficient_trace_data)."
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — TIERED EVALUATION STRATEGY
# ══════════════════════════════════════════════════════════════════════════════

heading1("Tiered Evaluation Strategy")

body(
    "The ten requirements are ordered by execution priority. Tier 1 checks gate all downstream "
    "evaluation — if either fails, the response is flagged and deeper checks are not required. "
    "Tier 2 runs on all responses that pass Tier 1. Tier 3 runs on flagged responses, high-value "
    "transactions, and periodic audit samples."
)

callout_body(
    "Regulated Industry Note: Financial services (MiFID II, Reg BI, SEC), healthcare (HIPAA, "
    "FDA clinical decision support regulations), legal, and other regulated industries impose "
    "requirements that elevate certain Tier 3 checks to mandatory status. REQ-10 (Source "
    "Traceability) is a compliance-mandatory audit trail requirement in these industries. "
    "REQ-09 (Confidence Calibration) is subject to suitability and appropriateness regulations. "
    "Full support for configurable risk-profile-based tier elevation requires the Regulated "
    "Industry Risk Profile Engine described in the Platform Roadmap. Until that enhancement is "
    "available, regulated deployments should treat Tier 3 as fully mandatory rather than on-demand."
)

# ── Tier table helper ─────────────────────────────────────────────────────────
def add_tier_table(rows_data, col3_label="Why It Gates/Why Standard/Why Deep"):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    tier_headers = ["Requirement", "Phase", "What It Catches", col3_label]
    hdr = table.rows[0].cells
    for i, h in enumerate(tier_headers):
        set_cell_bg(hdr[i], COLOR_HEADER)
        set_cell_borders(hdr[i], "1B2A4A", "6")
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = COLOR_WHITE

    for idx, (req, phase, catches, why) in enumerate(rows_data):
        row = table.add_row()
        cells = row.cells
        bg = COLOR_ROW1 if idx % 2 == 0 else COLOR_WHITE
        for ci, text in enumerate([req, phase, catches, why]):
            set_cell_bg(cells[ci], bg)
            set_cell_borders(cells[ci])
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cells[ci].paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_BODY
            if ci == 0:
                run.bold = True
                run.font.color.rgb = COLOR_TEAL
            if ci == 1:
                run.font.color.rgb = COLOR_SLATE

    for row in table.rows:
        row.cells[0].width = Inches(1.4)
        row.cells[1].width = Inches(1.3)
        row.cells[2].width = Inches(2.0)
        row.cells[3].width = Inches(1.7)

    doc.add_paragraph()


# ── Tier 1 ────────────────────────────────────────────────────────────────────
heading2("Tier 1 — Gate  (Run on every evaluation)")

tier1_rows = [
    (
        "REQ-01\nAction Verification",
        "Phase 1 — Available Now",
        "Every claimed action must map to a tool invocation span with a non-empty successful "
        "output. Fabricated confirmations and unverified actions caught here.",
        "If the claimed action never happened or returned nothing the entire response is built "
        "on a fabricated foundation. No further check can redeem it.",
    ),
    (
        "REQ-02\nTool Output Faithfulness",
        "Phase 1 — Available Now",
        "The agent's response must accurately reflect what tools returned. Overrides, "
        "substitutions, and misrepresentations caught here.",
        "If the action happened but was misreported the user receives false information about "
        "a real event. All downstream entity and accuracy checks become meaningless against "
        "a falsified tool result.",
    ),
]
add_tier_table(tier1_rows, "Why It Gates")

# ── Tier 2 ────────────────────────────────────────────────────────────────────
heading2("Tier 2 — Standard  (Run on all responses that pass Tier 1)")

tier2_rows = [
    (
        "REQ-03\nEntity Accuracy",
        "Phase 1 — Available Now",
        "All named entities must match source tool outputs or context. Substitutions, "
        "approximate matches, and invented entities each classified and flagged.",
        "Entity errors are the most operationally damaging hallucination class. A wrong "
        "identifier, location, or value causes real-world consequences that are often irreversible.",
    ),
    (
        "REQ-04\nUncertainty Acknowledgment",
        "Phase 1 — Available Now",
        "When a tool returns empty or non-parseable output the agent must flag uncertainty "
        "rather than fabricate a confident outcome.",
        "Fabrication from empty tool results is a primary failure mode in production agentic "
        "systems. Detecting it here prevents false confirmations reaching users.",
    ),
    (
        "REQ-05\nFactual Accuracy",
        "Phase 1 — Available Now",
        "Every factual claim must be traceable to source context or verified tool output.",
        "Even when tool outputs are faithfully reported agents can introduce unsupported facts "
        "in surrounding language summaries or inferences.",
    ),
    (
        "REQ-06\nReasoning Consistency",
        "Phase 1 — Available Now",
        "Sub-agent conclusions must not be materially altered at any agent-to-agent handoff "
        "boundary without a new tool result to justify the change. Produces no signal on "
        "single-agent deployments (trivially passes — no handoff boundaries present).",
        "The current evaluation template has no parameter to detect agent topology, so "
        "conditional promotion is not computable within the template. Elevated to Tier 2 "
        "for all agents. Handoff corruption is a primary failure mode in multi-agent "
        "systems and each boundary is an opportunity to introduce errors invisible in "
        "the final response alone.",
    ),
]
add_tier_table(tier2_rows, "Why Standard")

# ── Tier 3 ────────────────────────────────────────────────────────────────────
heading2("Tier 3 — Deep  (Flagged sessions, high-value transactions, periodic audits — mandatory for regulated deployments)")

tier3_rows = [
    (
        "REQ-07\nMulti-Turn Consistency",
        "Phase 1 — Available Now",
        "Facts established in earlier turns must not be contradicted in later turns without a "
        "corresponding tool result or context update.",
        "Long-running sessions accumulate state. Cross-turn contradictions are invisible to "
        "single-response evaluation but compound into major failures.",
    ),
    (
        "REQ-08\nScope Honesty",
        "Phase 1 — Available Now",
        "Tool inputs must preserve the specificity of the user's request. Input specificity "
        "drift — where the agent drops a qualifier or disambiguation element — flagged.",
        "Observable input drift causes tools to return results for the wrong target. Catching "
        "it here prevents silently scoped results from reaching users.",
    ),
    (
        "REQ-09\nConfidence Calibration",
        "Phase 1 — Available Now",
        "Agent's expressed certainty must match the strength of its tool evidence. Definitive "
        "language against empty or mismatched results flagged.",
        "False certainty is as harmful as false facts. Regulated: suitability and "
        "appropriateness requirements in financial services and clinical appropriateness "
        "standards in healthcare make this a compliance-relevant check.",
    ),
    (
        "REQ-10\nSource Traceability",
        "Phase 1 — Available Now",
        "Every entity in the agent's response must be traceable to a specific source span "
        "in the trace. Entities with no source span flagged as fabricated.",
        "Full audit capability. Regulated: source traceability is a compliance-mandatory "
        "requirement in financial services and healthcare — treat as Tier 2 in regulated "
        "deployments until Risk Profile Engine is available.",
    ),
]
add_tier_table(tier3_rows, "Why Deep")

# ── Automated Tier 3 Escalation ───────────────────────────────────────────────
heading2("Automated Tier 3 Escalation Rules")

body(
    "Because there is currently no mechanism for operators to specify which tier to execute "
    "per trace in Okahu, the evaluation system shall determine Tier 3 execution automatically "
    "using the rules below. Rules are derived from two signal sources: the outcome of each "
    "Tier 2 check, and trace metadata observable from span fields without human input. Rules "
    "are additive — the system collects all Tier 3 checks triggered across all rules and "
    "executes each check exactly once per trace regardless of how many rules triggered it."
)

# build escalation table
esc_table = doc.add_table(rows=1, cols=4)
esc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
esc_table.style = "Table Grid"
esc_headers = ["Rule", "Trigger Source", "Escalation Condition", "Tier 3 Checks Triggered"]
for i, h in enumerate(esc_headers):
    cell = esc_table.rows[0].cells[i]
    set_cell_bg(cell, COLOR_HEADER)
    set_cell_borders(cell, "1B2A4A", "6")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = COLOR_WHITE

esc_rows = [
    (
        "E-01",
        "Trace metadata",
        "scope.agentic.turn count > 1 detected in span fields — session has multiple turns.",
        "REQ-07 Multi-Turn Consistency",
    ),
    (
        "E-02",
        "REQ-03\nEntity Accuracy",
        "Result is minor_hallucination or major_hallucination — one or more entities deviate from source.",
        "REQ-10 Source Traceability",
    ),
    (
        "E-03",
        "REQ-03\nEntity Accuracy",
        "Result is major_hallucination — entity has no counterpart in any tool output or input span.",
        "REQ-08 Scope Honesty (Phase 1)",
    ),
    (
        "E-04",
        "REQ-04\nUncertainty Acknowledgment",
        "Result is major_hallucination — agent produced a confident claim from an empty tool output.",
        "REQ-09 Confidence Calibration",
    ),
    (
        "E-05",
        "REQ-05\nFactual Accuracy",
        "Result is minor_hallucination or major_hallucination — one or more claims lack a traceable source.",
        "REQ-10 Source Traceability",
    ),
    (
        "E-06",
        "REQ-06\nReasoning Consistency",
        "Result is major_hallucination — a sub-agent conclusion was materially altered at a handoff boundary.",
        "REQ-07 Multi-Turn Consistency\nREQ-08 Scope Honesty (Phase 1)",
    ),
    (
        "E-07",
        "Any Tier 1 or Tier 2 check",
        "Any check returns major_hallucination — a confirmed major failure at any earlier tier.",
        "REQ-09 Confidence Calibration\nREQ-10 Source Traceability",
    ),
]

for idx, (rule, trigger, condition, triggered) in enumerate(esc_rows):
    row = esc_table.add_row()
    cells = row.cells
    bg = COLOR_ROW1 if idx % 2 == 0 else COLOR_WHITE
    for ci, text in enumerate([rule, trigger, condition, triggered]):
        set_cell_bg(cells[ci], bg)
        set_cell_borders(cells[ci])
        cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cells[ci].paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_BODY
        if ci == 0:
            run.bold = True
            run.font.color.rgb = COLOR_TEAL
        if ci == 1:
            run.font.color.rgb = COLOR_SLATE

for row in esc_table.rows:
    row.cells[0].width = Inches(0.6)
    row.cells[1].width = Inches(1.4)
    row.cells[2].width = Inches(2.6)
    row.cells[3].width = Inches(1.8)

doc.add_paragraph()

body(
    "Default — No Escalation: If all Tier 1 and Tier 2 checks return no_hallucination and "
    "the trace contains a single turn (scope.agentic.turn = 1), Tier 3 is not executed for "
    "that trace. This is the expected outcome for a correctly functioning agent on a "
    "straightforward single-turn request."
)

body(
    "De-duplication: Rules E-02 and E-07 both trigger REQ-10; Rules E-01 and E-06 both "
    "trigger REQ-07. When multiple rules trigger the same Tier 3 check, that check is "
    "executed once. The evaluation system shall log which rules triggered each Tier 3 "
    "execution so that the escalation path is auditable."
)

callout_body(
    "Platform Note: These rules are computable entirely from Tier 1/2 outcomes and span "
    "metadata — no operator configuration is required. Implementation requires the evaluation "
    "engine to pass Tier 1/2 outcomes and trace metadata into a rule evaluation step before "
    "determining which Tier 3 checks to enqueue. The Regulated Industry Risk Profile Engine "
    "(Phase 2) will extend this by allowing operators to configure additional escalation rules "
    "and override the default no-escalation condition for high-risk deployments."
)

body(
    "REQ-01 and REQ-02 together would have flagged all six LG Travel Agent traces that were "
    "incorrectly labelled no_hallucination — four via REQ-01 (empty tool output with fabricated "
    "confirmation) and all six via REQ-02 (tool result overridden or fabricated in the agent "
    "response). Under the automated escalation rules, E-07 would have additionally triggered "
    "REQ-09 and REQ-10 for all six traces, and E-01 would have triggered REQ-07 for any "
    "multi-turn traces in the set. See Appendix for trace detail."
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — INDIVIDUAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════

heading1("Individual Requirements")

# ── REQ-01 ────────────────────────────────────────────────────────────────────
req_id_para("REQ-01", "Action Verification — Tier 1 Gate — Phase 1")

field_label("Requirement")
body(
    "The evaluation system shall verify that every consequential action the agent claims to "
    "have completed — including API mutations, database writes, external communications, "
    "resource allocations, and domain-specific transactions as defined in the agent's action "
    "registry — has a corresponding tool invocation span in the trace with a non-empty, "
    "successful output. An agent claiming a completed action with no supporting tool "
    "invocation, or with a tool invocation that returned an empty or error response, shall "
    "be flagged as a major hallucination."
)

field_label("Rationale")
body(
    "Agents can confirm actions they never took or that failed silently. In agentic systems "
    "where tool execution is asynchronous or logged separately, the only ground truth for "
    "whether an action occurred is the tool span record. Claims not backed by a successful "
    "tool span are unverifiable at best and fabricated at worst. This failure mode occurs "
    "across every industry and agent type."
)

field_label("Acceptance Criteria")
bullet("The evaluator shall map every action claimed in the agent's final response to a specific agentic.tool.invocation span in the trace.")
bullet("Each span shall have a status_code of OK and a non-empty data.output.")
bullet("Action claims with no matching span shall be flagged as major hallucinations.")
bullet("Action claims where the matching span returned an empty output ({}, null, or error) shall be flagged as major hallucinations.")
bullet("Action claims where the matching span returned a success but for a different target or value than claimed shall be flagged as major hallucinations.")

field_label("Regulated Industry Note")
body(
    "This check carries critical severity in all industries. In healthcare, a claimed "
    "medication administration or procedure confirmation that did not occur is a patient "
    "safety incident. In financial services, a claimed trade execution or fund transfer "
    "that did not occur is a financial integrity failure. In both cases the consequence "
    "extends beyond product quality into regulatory and liability exposure."
)

field_label("Examples")
add_examples_table([
    ("no_hallucination",
     "Agent claims: 'Your prescription has been sent to the pharmacy.' Matching "
     "agentic.tool.invocation span for send_prescription exists with status OK and output "
     "confirming successful transmission to the correct pharmacy. Action is fully verified."),
    ("minor_hallucination",
     "Agent claims: 'Your transaction of $250.00 has been processed.' Matching tool span "
     "exists with status OK but output confirms the transaction as $245.00. The action "
     "occurred but for a different amount than claimed."),
    ("major_hallucination",
     "Agent claims: 'Your insurance claim CLM-8821 has been submitted.' Matching "
     "agentic.tool.invocation span exists but data.output is {}. No successful action is "
     "evidenced — the confirmation is fabricated."),
    ("major_hallucination",
     "LG Travel Agent (trace 182c5571): Hotel tool returned {}. Lodging agent responded: "
     "'Hotel in Paris, TX booked April 17\u201320.' The agentic.tool.invocation span for the "
     "hotel booking has an empty data.output — no successful booking action is evidenced. "
     "The confirmation is fabricated."),
    ("unknown_not_computable",
     "The current evaluation template has no parameter to map claimed actions to specific "
     "agentic.tool.invocation spans. Whether any action claim has a corresponding verified "
     "tool execution is uncomputable under the current template."),
])

# ── REQ-02 ────────────────────────────────────────────────────────────────────
req_id_para("REQ-02", "Tool Output Faithfulness — Tier 1 Gate — Phase 1")

field_label("Requirement")
body(
    "The evaluation system shall verify that the agent's response accurately reflects the "
    "literal content of every tool output received during the session. The agent shall not "
    "paraphrase, correct, substitute, or omit material facts present in tool outputs when "
    "reporting results to the user. This applies at every agent-to-agent handoff boundary "
    "in the system topology — whether supervisor-to-sub-agent, sequential chain, or "
    "parallel coordination."
)

field_label("Rationale")
body(
    "In multi-agent systems, each relay between agents is an opportunity to alter, soften, "
    "or override what a tool actually returned. A faithful evaluation must compare the "
    "tool's raw output against what each agent communicated, at every point in the chain."
)

field_label("Acceptance Criteria")
bullet("The evaluator shall extract all agentic.tool.invocation span data.output events from the trace.")
bullet("The evaluator shall compare each tool output to the agent's corresponding response at the nearest downstream turn_end span.")
bullet("Any deviation — substitution, omission of material fact, invented detail, or contradiction — shall be classified and flagged.")
bullet("If a tool output is empty or malformed, the evaluator shall record that condition before assessing faithfulness.")
bullet("The evaluator shall check faithfulness at each agent handoff boundary, not only in the final response.")

field_label("Regulated Industry Note")
body(
    "In financial services, misrepresenting a tool output — such as reporting a trade as "
    "executed at a different price, or a fund balance as a different amount — may constitute "
    "a regulatory violation independent of intent. In healthcare, misrepresenting a clinical "
    "tool result — such as a diagnostic value or medication interaction flag — is a patient "
    "safety issue. Tool output faithfulness must be treated as a compliance check in these contexts."
)

field_label("Examples")
add_examples_table([
    ("no_hallucination",
     "Tool returned: 'Appointment confirmed — Dr. Chen, April 20 at 2:00 PM, Cardiology.' "
     "Agent response: 'Your appointment with Dr. Chen in Cardiology on April 20 at 2:00 PM "
     "is confirmed.' Faithful relay with acceptable standard formatting."),
    ("minor_hallucination",
     "Tool returned: 'Account balance: $4,231.50 as of April 20, 2026.' Agent response: "
     "'Your account balance is approximately $4,200.' The agent rounded the figure and "
     "dropped the as-of date qualifier."),
    ("major_hallucination",
     "Tool returned: 'Claim denied — policy exclusion applies to pre-existing conditions.' "
     "Agent response: 'Your claim has been approved and payment will be issued within 5 "
     "business days.' The tool result was overridden with the opposite outcome."),
    ("major_hallucination",
     "LG Travel Agent (trace df7648b4): Hotel tool returned 'H\u00f4tel R\u00e9publique "
     "(Paris, France).' Lodging agent response: 'Your hotel in Paris, Texas has been "
     "successfully booked.' The agent substituted the tool's location (Paris, France) with "
     "the user's desired location (Paris, Texas) — a direct override of the tool output."),
    ("unknown_not_computable",
     "The current evaluation template has no tool_output_faithfulness parameter. It has no "
     "field to compare agentic.tool.invocation outputs against agent responses. All tool "
     "faithfulness assessments — whether the agent accurately relayed, altered, or fabricated "
     "a tool result — are uncomputable under the current template."),
])

# ── REQ-03 ────────────────────────────────────────────────────────────────────
req_id_para("REQ-03", "Entity Accuracy — Tier 2 Standard")

field_label("Requirement")
body(
    "The evaluation system shall verify that all named entities in the agent's final response "
    "exactly match the entities present in the source tool outputs or input context. Entity "
    "types include but are not limited to: identifiers, locations, monetary values, clinical "
    "terms, dates, codes, and proper names. Substitutions, approximate matches, and invented "
    "entities shall each be classified and flagged."
)

field_label("Rationale")
body(
    "Entity-level errors are the most operationally damaging class of hallucination. A wrong "
    "patient identifier, account number, drug name, or location causes real-world harm that "
    "is often irreversible. Claim-level accuracy checks are insufficient — an agent can make "
    "a true claim while embedding a false entity within it."
)

field_label("Acceptance Criteria")
bullet("The evaluator shall extract every named entity from the agent's final response.")
bullet("For each entity, the evaluator shall identify its counterpart in the tool output or input context.")
bullet("Standardised transformations — capitalisation normalisation, accepted domain abbreviations, and well-known code expansions — shall be classified as no hallucination as they are lossless and unambiguous.")
bullet("Entities that differ from source in a small but factually meaningful way — dropped accent on proper noun, inferred unit not in source, computed value from source data — shall be classified as minor hallucinations.")
bullet("Entities that differ from source in substance — wrong identifier, wrong location, wrong name, wrong value — shall be classified as major hallucinations.")
bullet("Entities with no counterpart in any source shall be classified as major hallucinations.")

field_label("Regulated Industry Note")
body(
    "In healthcare, entity errors at the patient identifier, medication name, dosage, or "
    "procedure code level are patient safety incidents regardless of whether the downstream "
    "action was taken. In financial services, entity errors at the account number, transaction "
    "amount, security identifier (CUSIP, ISIN, ticker), or regulatory reference level carry "
    "potential compliance and financial liability. See the Entity Schema Registry in the "
    "Phase 2 Platform Roadmap for domain-specific entity type and transformation rule "
    "configuration required for compliance-grade entity evaluation."
)

field_label("Examples")
add_examples_table([
    ("no_hallucination",
     "Tool returned patient record for ID PAT-00482. Agent response references "
     "'patient PAT-00482.' Exact entity match."),
    ("minor_hallucination",
     "LG Travel Agent (trace df7648b4): Tool returned hotel name 'H\u00f4tel R\u00e9publique.' "
     "Agent response stated 'Hotel Republique.' The accent was dropped from a proper noun "
     "present in the tool output — a small but real alteration to a specific named entity."),
    ("major_hallucination",
     "User requested account information for ACC-7821. Agent response provides details for "
     "ACC-7812 — a transposed digit. The entity is wrong and could expose another user's "
     "financial data."),
    ("major_hallucination",
     "LG Travel Agent (trace df7648b4): Tool returned 'H\u00f4tel R\u00e9publique (Paris, "
     "France).' Agent response confirmed a booking for 'a hotel in Paris, Texas.' The "
     "location entity was substituted — Paris, France \u2192 Paris, Texas — a direct major "
     "entity substitution with real-world booking consequences."),
    ("unknown_not_computable",
     "The current evaluation template has no parameter to enumerate individual named entities "
     "and trace each to a source span. Its factual_alignments field operates at the claim "
     "level against context documents — not at the entity level against span IDs. Per-entity "
     "source verification is uncomputable under the current template."),
])

# ── REQ-04 ────────────────────────────────────────────────────────────────────
req_id_para("REQ-04", "Uncertainty Acknowledgment — Tier 2 Standard")

field_label("Requirement")
body(
    "The evaluation system shall verify that when a tool returns an empty, malformed, or "
    "non-parseable result, the agent acknowledges the uncertainty rather than fabricating a "
    "confident outcome. An agent that proceeds as if a tool succeeded when the tool returned "
    "no usable data shall be flagged."
)

field_label("Rationale")
body(
    "Empty or error tool responses are a primary trigger for fabrication in agentic systems. "
    "An agent that cannot distinguish between a successful tool result and an empty one will "
    "fill the gap with plausible-sounding invented content, presenting it with the same "
    "confidence as a real result."
)

field_label("Acceptance Criteria")
bullet("The evaluator shall inspect the data.output of every agentic.tool.invocation span for empty, null, or non-parseable payloads.")
bullet("If a tool returned {}, null, or an error, the evaluator shall check whether the agent's subsequent response appropriately hedges or flags the uncertainty.")
bullet("An agent that produces specific, confident claims following an empty tool response shall be classified as a major hallucination.")
bullet("An agent that produces vague but non-committal language following an empty tool response shall be classified as a minor hallucination.")

field_label("Regulated Industry Note")
body(
    "In healthcare, proceeding without acknowledging insufficient information is a patient "
    "safety issue — not a quality issue. A clinical decision support agent that produces a "
    "recommendation when its underlying data retrieval returned empty is operating outside "
    "safe parameters. In financial services, making a recommendation or executing a "
    "transaction when data retrieval failed without flagging the gap may violate suitability "
    "requirements. The threshold for what constitutes adequate uncertainty acknowledgment "
    "should be configured more strictly for these deployments."
)

field_label("Examples")
add_examples_table([
    ("no_hallucination",
     "Tool returns {}. Agent response: 'I was unable to confirm your booking — the system "
     "returned no confirmation details. Please try again or contact support.' Uncertainty "
     "correctly acknowledged."),
    ("minor_hallucination",
     "Tool returned a valid result but included a data quality warning flag the agent did not "
     "relay to the user. Agent confirmed the action without noting the caveat present in the "
     "tool output."),
    ("major_hallucination",
     "Tool returned {}. Agent response: 'Your appointment has been confirmed for April 20 at "
     "2:00 PM with Dr. Chen.' Complete appointment details fabricated from an empty tool response."),
    ("major_hallucination",
     "LG Travel Agent (trace 182c5571): Hotel tool returned {}. Lodging agent responded: "
     "'Hotel in Paris, TX booked April 17\u201320.' Specific dates and location fabricated "
     "with full confidence from an empty tool response. This failure pattern was observed "
     "across multiple April 10 traces where the hotel tool consistently returned {}."),
    ("unknown_not_computable",
     "The current evaluation template has no parameter to inspect the tool's data.output "
     "payload before evaluating the agent's response. Uncertainty acknowledgment assessment "
     "is uncomputable under the current template."),
])

# ── REQ-05 ────────────────────────────────────────────────────────────────────
req_id_para("REQ-05", "Factual Accuracy — Tier 2 Standard — Phase 1")

field_label("Requirement")
body(
    "The evaluation system shall compare every factual claim in the agent's final response "
    "against the source context provided in the prompt or the verified output of tool "
    "invocations. Any claim that cannot be traced to either source shall be flagged."
)

field_label("Rationale")
body(
    "Agents operating on retrieved or tool-generated data can introduce incorrect facts when "
    "paraphrasing, summarising, or completing partial information. Without explicit tracing of "
    "each claim back to its source, factual errors pass undetected — even when the tool "
    "execution itself was successful and faithfully reported."
)

field_label("Acceptance Criteria")
bullet("The evaluator shall identify each distinct factual claim in the agent response.")
bullet("For each claim, the evaluator shall locate the corresponding statement in the input context or tool span data.output.")
bullet("Claims with no traceable source shall be flagged.")
bullet("The evaluator shall distinguish between claims sourced from tool outputs, claims sourced from context documents, and claims that are agent-generated inferences.")
bullet("The evaluator shall list which claims were supported and which were not.")

field_label("Regulated Industry Note")
body(
    "In financial services, incorrect figures — interest rates, account balances, transaction "
    "amounts, regulatory thresholds — have direct monetary impact and may constitute misleading "
    "communication under securities law. In healthcare, incorrect clinical facts — dosages, "
    "contraindications, diagnostic criteria — carry direct patient safety risk. For these "
    "deployments, the factual accuracy check should be configured with zero tolerance for "
    "unsupported claims in the domain entities defined in the entity schema."
)

field_label("Examples")
add_examples_table([
    ("no_hallucination",
     "Tool returned account balance of $4,231.50. Agent states: 'Your current account "
     "balance is $4,231.50.' Claim is fully traceable to the tool output."),
    ("minor_hallucination",
     "Tool returned temperature of 72 with no unit specified. Agent states: '72 degrees "
     "Fahrenheit.' The unit is an inference added by the agent — not present in the tool "
     "output. Reasonable in context but an unverifiable addition."),
    ("major_hallucination",
     "Tool returned: 'Claim status: denied.' Agent states: 'Your claim has been approved "
     "and payment of $1,200 will be issued within 5 business days.' Multiple fabricated "
     "facts — approval status, payment amount, timeline — none present in the tool output."),
    ("major_hallucination",
     "LG Travel Agent (trace 182c5571): Hotel tool returned {}. Agent stated: 'Hotel in "
     "Paris, TX booked April 17\u201320.' The check-in and check-out dates (April 17\u201320) "
     "are factual claims that appear nowhere in the tool output span. They are agent-generated "
     "fabrications presented as confirmed booking facts."),
    ("unknown_not_computable",
     "The current evaluation template has no parameter to check tool output payloads as a "
     "source for factual claims. Its factual_alignments field checks claims against context "
     "documents only. Whether a claim originates from a tool result vs. fabrication is "
     "uncomputable under the current template."),
])

# ── REQ-06 ────────────────────────────────────────────────────────────────────
req_id_para("REQ-06", "Reasoning Consistency — Tier 2 Standard — Phase 1")

field_label("Requirement")
body(
    "The evaluation system shall verify that the agent's stated reasoning is internally "
    "consistent with its final conclusion, and that conclusions produced at any agent-to-agent "
    "handoff boundary in the system topology are not materially altered or contradicted by the "
    "receiving agent without a new tool result or context update to justify the change."
)

field_label("Tier Rationale")
body(
    "This check is elevated to Tier 2 Standard for all agent deployments. The current "
    "evaluation template has no parameter to detect agent topology — there is no field "
    "for number of agents, handoff boundaries, or span types such as delegation or "
    "entity.from_agent. Conditional promotion (Tier 2 for multi-agent, Tier 3 for "
    "single-agent) is therefore not computable within the current template. On single-agent "
    "deployments this check produces no signal and trivially passes. On multi-agent "
    "deployments — which represent the primary production pattern for agentic systems — "
    "handoff corruption is a primary failure mode that is invisible to single-response "
    "evaluation and not caught by Tier 1 alone."
)

field_label("Rationale")
body(
    "In multi-agent architectures — whether supervisor-based, sequential chain, parallel, or "
    "peer-to-peer — each handoff is an opportunity to reinterpret or override what the prior "
    "agent established. Inconsistencies between what an agent concluded and what the next "
    "agent relayed indicate that the receiving agent is reasoning from something other than "
    "the actual results it received."
)

field_label("Acceptance Criteria")
bullet("The evaluator shall extract all turn_end outputs from each agent within the same trace, in topology order.")
bullet("The evaluator shall compare each agent's stated result against the downstream agent's summary or relay of that result.")
bullet("Any contradiction — different outcome, different entity, different conclusion — introduced without a new tool call shall be flagged.")
bullet("The evaluator shall identify which agent in the chain introduced the inconsistency.")
bullet("The evaluator shall distinguish between complete contradiction (major) and partial omission of qualifying detail (minor).")
bullet("On single-agent deployments where no handoff boundaries are present, the evaluator shall return no_hallucination without further processing.")

field_label("Regulated Industry Note")
body(
    "Regulated industries require explainable and auditable reasoning chains. An alteration "
    "in the relay between agents is not only a quality issue — in financial services it may "
    "represent a material misrepresentation of a recommendation basis, and in healthcare it "
    "may break the clinical decision chain in a way that is invisible to the end user but "
    "consequential for care."
)

field_label("Examples")
add_examples_table([
    ("no_hallucination",
     "Sub-agent concludes: 'Patient eligibility verified — all criteria met for this "
     "procedure.' Supervisor relay: 'Patient eligibility for the procedure has been "
     "confirmed.' Consistent at the handoff boundary."),
    ("minor_hallucination",
     "Not observed in the LG Travel Agent reference traces. Reasoning consistency failures "
     "in those traces were binary — either fully consistent or directly contradicted. In "
     "broader agentic systems this category does occur. Example: a billing sub-agent reasons "
     "'Customer was charged twice for order #8821 due to a system timeout — refund of $47.99 "
     "approved under the duplicate charge policy.' The supervisor relays: 'A refund of $47.99 "
     "has been approved for order #8821.' The conclusion is preserved but the causal reasoning "
     "— duplicate charge, system timeout, policy basis — was stripped in the relay. Downstream "
     "systems relying on a reason code to process the refund correctly would be affected."),
    ("major_hallucination",
     "Sub-agent concludes: 'Loan application does not meet minimum credit score requirements "
     "— declined.' Supervisor relay: 'Loan application is under review.' The conclusion was "
     "changed from a definitive decline to an ambiguous open state at the handoff boundary."),
    ("major_hallucination",
     "LG Travel Agent (trace 7263beb5): Hotel tool returned {}. Lodging sub-agent output "
     "reflected uncertainty about hotel availability. Supervisor relay at the handoff "
     "boundary: 'Hotel de la Seine in Paris, Tennessee has been reserved for your stay.' "
     "The sub-agent's uncertainty was overridden with a specific fabricated hotel name and "
     "confirmed booking — a contradiction introduced at the handoff with no new tool result."),
    ("unknown_not_computable",
     "The current evaluation template has no parameter to compare sub-agent turn_end outputs "
     "against downstream agent summaries. Cross-agent reasoning consistency is uncomputable "
     "under the current template."),
])

# ── REQ-07 ────────────────────────────────────────────────────────────────────
req_id_para("REQ-07", "Multi-Turn Consistency — Tier 3 Deep — Phase 1")

field_label("Requirement")
body(
    "The evaluation system shall verify that factual claims established in earlier turns or "
    "earlier agent invocations within the same session are not contradicted in later turns "
    "without a corresponding tool result or context update that justifies the change."
)

field_label("Rationale")
body(
    "Multi-step agentic sessions accumulate state across turns. Each agent invocation has "
    "access to prior context and can silently alter established facts. A consistent agent "
    "must either preserve prior facts or explicitly acknowledge when a fact has changed and why."
)

field_label("Acceptance Criteria")
bullet("The evaluator shall extract all data.input and data.output events across all spans in the trace in chronological order using scope.agentic.turn and scope.agentic.invocation fields.")
bullet("The evaluator shall identify facts established in early turns — confirmed outcomes, entities, identifiers, and values.")
bullet("The evaluator shall check whether those facts appear consistently in all subsequent turns and in the final response.")
bullet("Any change to an established fact not supported by a new tool result or explicit context update shall be flagged.")

field_label("Regulated Industry Note")
body(
    "In financial services, suitability assessments (Reg BI, MiFID II) made across a session "
    "must remain consistent — a product assessed as high-risk in one turn must not be "
    "recommended as suitable for a conservative investor in a later turn without documented "
    "justification. In healthcare, clinical decisions made across a patient interaction must "
    "not contradict without explicit clinical reasoning. Until the Regulated Industry Risk "
    "Profile Engine is available, regulated deployments should treat this check as Tier 2."
)

field_label("Examples")
add_examples_table([
    ("no_hallucination",
     "Turn 1: Agent confirms patient allergy to penicillin-class antibiotics. Turn 4: Agent "
     "recommends an alternative antibiotic, correctly excluding penicillin-class drugs. Prior "
     "established fact is preserved across turns."),
    ("minor_hallucination",
     "Turn 2: Agent confirms order total as $847.50 including all fees. Turn 5: Supervisor "
     "summarises 'your order' without restating the total — the established figure was "
     "weakened but not contradicted."),
    ("major_hallucination",
     "Turn 2: Agent classifies an investment product as high-risk based on volatility "
     "analysis. Turn 6: Agent recommends the same product as 'suitable for conservative "
     "portfolios.' Direct contradiction of an established risk classification across turns "
     "with no new analysis or tool result to justify the change."),
    ("major_hallucination",
     "LG Travel Agent (trace df7648b4): In an earlier turn the hotel sub-agent returned "
     "'H\u00f4tel R\u00e9publique (Paris, France)' — establishing Paris, France as the hotel "
     "location within the session. In the final supervisor turn: 'Your hotel in Paris, Texas "
     "has been successfully booked.' The location entity (France) was contradicted in a later "
     "turn with no new tool result or context update to justify the change."),
    ("unknown_not_computable",
     "The current evaluation template evaluates a single response in isolation. It has no "
     "parameter to track facts established in prior turns or prior agent invocations within "
     "the same session. Cross-turn consistency checks are uncomputable under the current template."),
])

# ── REQ-08 ────────────────────────────────────────────────────────────────────
req_id_para("REQ-08", "Scope Honesty — Tier 3 Deep")

field_label("Requirement")
body(
    "The evaluation system shall verify that the agent does not present results for inputs "
    "that are outside its operational scope, and that when inputs are ambiguous the agent "
    "either resolves the ambiguity explicitly or flags it rather than proceeding with a "
    "silent assumption."
)

field_label("Rationale")
body(
    "Agents operating with ambiguous inputs may silently resolve ambiguity in ways that "
    "produce results for the wrong target. Presenting those results without flagging the "
    "ambiguity constitutes a scope violation — the agent acted beyond what its inputs could "
    "reliably support. In regulated industries, scope violations may additionally constitute "
    "unauthorized data access."
)

field_label("Acceptance Criteria")
bullet("The evaluator shall compare tool input parameters against the specificity of the user's original request as recorded in the earliest data.input event in the session.")
bullet("Any case where the tool was called with less specific inputs than the user provided — such as dropping a state qualifier, account suffix, or product variant — shall be flagged.")
bullet("Any case where the tool returned results for a different scope than requested without the agent noting the discrepancy shall be flagged.")
bullet("The evaluator shall check whether the agent explicitly resolved or acknowledged ambiguity.")

field_label("Regulated Industry Note")
body(
    "In regulated industries, scope authorization is a compliance requirement — not a "
    "quality check. An agent accessing a patient record outside its clinical authorization "
    "(HIPAA), or a financial agent querying client data without appropriate entitlement "
    "(MiFID II data access controls, SEC Reg S-P), represents a regulatory violation "
    "regardless of whether the result was accurate. This check (Phase 1) detects input "
    "specificity drift from span data. Compliance-grade scope authorization comparison "
    "against declared agent boundaries is available in Phase 2 via the Authorized Scope "
    "Registry — see the Phase 2 Platform Roadmap. Until that enhancement is available, "
    "regulated deployments must implement scope authorization controls at the platform "
    "infrastructure level and treat this check as Tier 2."
)

field_label("Examples")
add_examples_table([
    ("no_hallucination",
     "User requests 'balance for savings account ending 4821.' Tool called with "
     "{account_type: 'savings', account_suffix: '4821'}. Agent response attributes result "
     "to 'savings account ending 4821' explicitly. Full specificity preserved."),
    ("minor_hallucination",
     "User requests 'pricing for Product SKU-7821-A.' Tool called with {sku: 'SKU-7821'} "
     "— variant suffix dropped. Agent returns pricing attributed to 'your product' without "
     "flagging that variant disambiguation was lost."),
    ("major_hallucination",
     "LG Travel Agent (traces df7648b4, 60a156dac): User requested a hotel in Paris, Texas. "
     "Hotel tool called with {city: 'Paris'} — state qualifier dropped. Tool returned results "
     "for Paris, France. Agent reported the France result as the user's Paris, Texas booking "
     "with no acknowledgment of the scope discrepancy. Observed across multiple April 9 traces."),
    ("unknown_not_computable",
     "The current evaluation template has no parameter to compare tool input parameters "
     "against the specificity of the user's original request. Scope drift between user "
     "intent and tool invocation is uncomputable under the current template."),
])

# ── REQ-09 ────────────────────────────────────────────────────────────────────
req_id_para("REQ-09", "Confidence Calibration — Tier 3 Deep — Phase 1")

field_label("Requirement")
body(
    "The evaluation system shall verify that the agent's expressed level of certainty is "
    "proportional to the strength of the evidence supporting its claims. An agent that "
    "expresses high confidence when its tool results are empty, ambiguous, or contradictory "
    "shall be flagged."
)

field_label("Rationale")
body(
    "False certainty is as harmful as false facts. Users make irreversible decisions — "
    "financial commitments, clinical actions, legal agreements — based on how confident an "
    "agent appears. The expressed confidence must match the actual evidentiary basis, not "
    "the plausibility of the claim."
)

field_label("Acceptance Criteria")
bullet("The evaluator shall assess the language of certainty in the agent's response: definitive completions ('has been confirmed,' 'is approved,' 'will be issued') vs. hedged language ('appears to have been,' 'I was unable to confirm,' 'based on available data').")
bullet("The evaluator shall compare certainty language against the quality of the underlying tool output: successful with full detail, successful with partial detail, empty, or error.")
bullet("Definitive language backed by empty or error tool outputs shall be classified as major hallucinations.")
bullet("Definitive language backed by mismatched tool outputs shall be classified as minor or major hallucinations depending on degree of mismatch.")
bullet("The evaluator shall distinguish between the agent's expressed confidence (this check) and the evaluator's own confidence in its assessment (a separate metadata field).")

field_label("Regulated Industry Note")
body(
    "Financial services suitability regulations (MiFID II suitability assessment, SEC Reg BI "
    "best interest standard) require that recommendations are appropriate to the client's "
    "circumstances and that uncertainty is disclosed. An agent expressing high confidence in "
    "a recommendation based on incomplete or ambiguous data may violate these requirements. "
    "In healthcare, clinical appropriateness standards require that the agent's confidence "
    "in a clinical recommendation matches its evidentiary basis. Until the Regulated Industry "
    "Risk Profile Engine is available, regulated deployments should treat this check as Tier 2."
)

field_label("Examples")
add_examples_table([
    ("no_hallucination",
     "Tool returned a complete verified record with all required fields populated and status "
     "confirmed. Agent response: 'Your application has been successfully submitted and "
     "confirmed.' High certainty language is fully supported."),
    ("minor_hallucination",
     "Tool returned a result but included a data quality warning flag. Agent confirmed the "
     "result with full certainty language without qualifying the flagged field."),
    ("major_hallucination",
     "Tool returned {}. Agent response: 'Your investment of $50,000 has been confirmed and "
     "will begin generating returns at the agreed rate immediately.' Maximum confidence "
     "language with zero evidentiary support from the tool."),
    ("major_hallucination",
     "LG Travel Agent (trace 182c5571): Hotel tool returned {}. Agent responded: 'Hotel in "
     "Paris, TX booked April 17\u201320.' The word 'booked' is maximum-certainty confirmation "
     "language — no hedging, no qualification — with an empty tool output as the sole "
     "evidentiary basis. This confidence level is completely unsupported."),
    ("unknown_not_computable",
     "The current evaluation template's confidence_level field measures the evaluator's "
     "confidence in its own assessment — not the agent's expressed certainty relative to "
     "the strength of its tool evidence. Whether the agent overclaimed or underclaimed "
     "relative to its actual tool results is uncomputable under the current template."),
])

# ── REQ-10 ────────────────────────────────────────────────────────────────────
req_id_para("REQ-10", "Source Traceability — Tier 3 Deep")

field_label("Requirement")
body(
    "The evaluation system shall verify that every specific entity in the agent's response "
    "can be traced to a specific span in the trace — identified by span ID — from which that "
    "entity originated. Entities that cannot be sourced to a specific span shall be flagged."
)

field_label("Rationale")
body(
    "Traceability requires not just that a claim is plausible but that it is sourced to a "
    "specific, identifiable point in the execution. This provides full audit capability and "
    "is the foundation for post-incident investigation, compliance reporting, and trust in "
    "the system's outputs."
)

field_label("Acceptance Criteria")
bullet("The evaluator shall enumerate all named entities in the agent's final response.")
bullet("For each entity, the evaluator shall identify the span ID of the span where that entity first appeared in the execution.")
bullet("The evaluator shall verify that the entity in the response matches the entity in the source span exactly.")
bullet("Entities with no traceable source span shall be flagged as fabricated.")
bullet("The evaluator shall produce a source map: entity → source span ID → span type for each entity in the response.")

field_label("Regulated Industry Note")
body(
    "Source traceability is a compliance-mandatory requirement in financial services and "
    "healthcare — not a quality check. MiFID II requires audit trails for investment "
    "recommendations. FDA clinical decision support regulations require traceability of "
    "clinical outputs to their data sources. HIPAA requires audit logs for protected health "
    "information access. For regulated deployments this check must be treated as Tier 2 "
    "mandatory, not Tier 3 on-demand. Domain-specific entity type definitions and "
    "transformation rules that extend Phase 1 traceability to compliance-grade coverage "
    "are available in Phase 2 via the Entity Schema Registry — see the Phase 2 Platform Roadmap."
)

field_label("Examples")
add_examples_table([
    ("no_hallucination",
     "Agent response references 'Patient Record PRN-00482, last updated April 10, 2026.' "
     "This exact reference appears in the data.output of the get_patient_record tool "
     "invocation span (span ID: a3b84a93). Full chain of custody from span to response "
     "is intact."),
    ("minor_hallucination",
     "Agent response references 'the Q3 report.' The tool output span referenced 'Q3 "
     "Financial Report, Draft v2, Internal Only.' The agent dropped the draft qualifier "
     "and the internal classification — a small but real omission from the source entity."),
    ("major_hallucination",
     "Tool returned {}. Agent response references 'booking confirmation BCK-9921.' No "
     "booking reference appears in any tool output span in the trace. The identifier is "
     "fabricated with no source span."),
    ("major_hallucination",
     "LG Travel Agent (trace 7263beb5): Hotel tool returned {}. Supervisor response "
     "references 'Hotel de la Seine in Paris, Tennessee.' The entity 'Hotel de la Seine' "
     "does not appear in any agentic.tool.invocation data.output span in the trace. "
     "It is a fabricated entity with no traceable source span — a complete chain-of-custody "
     "break from tool execution to final response."),
    ("unknown_not_computable",
     "The current evaluation template has no parameter to enumerate individual named entities "
     "and trace each to a source span ID. Its factual_alignments field operates at the claim "
     "level against context documents — not at the entity level against specific span IDs. "
     "Per-entity source tracing is uncomputable under the current template."),
])

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — PLATFORM ROADMAP: PHASE 2 ENHANCEMENTS
# ══════════════════════════════════════════════════════════════════════════════

heading1("Platform Roadmap — Phase 2 Enhancements")

body(
    "The following platform capabilities are required to fully implement certain evaluation "
    "requirements and to support configurable regulatory risk profiles. These are documented "
    "here to capture the requirements and their dependencies. They are not available in the "
    "current platform release."
)

# ── Enhancement 1 — Tool Contract Registry ───────────────────────────────────
heading2("Tool Contract Registry")

field_label("Description")
body(
    "A registry where operators declare the expected response schema for each tool registered "
    "in the platform, including: required fields in a successful response, valid empty states "
    "and what they signify, error response signatures, and the distinction between a tool that "
    "legitimately returns no data vs. one that has failed silently."
)

field_label("Enables")
bullet("REQ-04 Phase 2: Full Uncertainty Acknowledgment evaluation — distinguishing valid empty responses from silent failures")
bullet("Eliminates unknown_not_computable (registry_not_configured) outcomes for REQ-04")
bullet("Provides richer context for REQ-01 Action Verification when tool outputs are ambiguous")

field_label("Platform Capabilities Required")
bullet("Tool contract schema definition interface for operators")
bullet("Contract storage and versioning per tool per agent deployment")
bullet("Evaluation engine integration to retrieve contract at evaluation time")
bullet("Contract validation — alerting operators when observed tool responses diverge from declared contract")

field_label("Evaluation Impact")
body(
    "Without this enhancement REQ-04 returns unknown_not_computable for all cases "
    "where a tool returned {} or a non-parseable response. With this enhancement the evaluator "
    "can make a determination in all cases."
)

field_label("REQ-04 Phase 2 Requirement Extension")
body(
    "When the Tool Contract Registry is available, the Uncertainty Acknowledgment check "
    "shall be extended as follows:"
)
bullet("The evaluator shall retrieve the declared response contract for each tool called during the trace.")
bullet("The evaluator shall determine whether {} or a non-parseable response is a declared valid empty state for that tool or represents a silent failure.")
bullet("If {} is a declared valid empty state and the agent correctly hedged, the evaluator shall classify the response as no_hallucination for this check.")
bullet("If {} is a declared valid empty state and the agent fabricated a confident outcome, the evaluator shall classify the response as major_hallucination.")
bullet("If the tool response diverges from its declared contract in structure or type (not just emptiness), the evaluator shall flag the contract violation as a separate finding alongside the hallucination classification.")
bullet("The evaluator shall no longer return unknown_not_computable (registry_not_configured) for REQ-04 once this enhancement is active.")

# ── Enhancement 2 — Authorized Scope Registry ────────────────────────────────
heading2("Authorized Scope Registry")

field_label("Description")
body(
    "A registry where operators declare the authorized operational boundaries for each agent "
    "in the system, including: permitted tools and APIs the agent may call, permitted data "
    "sources and record types the agent may access, permitted actions the agent may take, and "
    "any conditional scope restrictions (e.g. agent may access patient records only for "
    "patients in its assigned care team)."
)

field_label("Enables")
bullet("REQ-08 Phase 2: Full Scope Honesty evaluation — detecting agents operating outside permitted boundaries")
bullet("Eliminates unknown_not_computable (registry_not_configured) outcomes for REQ-08 Phase 2")
bullet("Compliance-grade scope violation detection for regulated industry deployments")

field_label("Platform Capabilities Required")
bullet("Scope declaration interface for operators with role-based access control")
bullet("Scope storage per agent per deployment with versioning")
bullet("Real-time scope comparison at evaluation time against observed tool calls in spans")
bullet("Scope violation alerting separate from hallucination classification")
bullet("Integration with identity and entitlement systems for regulated deployments")

field_label("Evaluation Impact")
body(
    "Without this enhancement REQ-08 is limited to input specificity drift detection "
    "from spans. With this enhancement the evaluator can detect unauthorized data access and "
    "out-of-scope actions — the compliance-critical case for regulated industries."
)

field_label("REQ-08 Phase 2 Requirement Extension")
body(
    "When the Authorized Scope Registry is available, the Scope Honesty check shall be "
    "extended as follows:"
)
bullet("The evaluator shall retrieve the declared authorized scope for each agent in the trace.")
bullet("The evaluator shall compare every tool call and data access observed in the trace spans against the declared scope for the agent that made the call.")
bullet("Any tool call, data source access, or action that falls outside the declared authorized scope shall be flagged as a scope violation, classified as major_hallucination, and reported as a separate finding.")
bullet("The evaluator shall identify which agent in the topology performed the out-of-scope action using entity.from_agent span metadata.")
bullet("Scope violations shall be reported with the specific span ID of the out-of-scope invocation to enable audit trail reconstruction.")
bullet("The evaluator shall no longer return unknown_not_computable (registry_not_configured) for the scope authorization check once this enhancement is active.")

# ── Enhancement 3 — Entity Schema Registry ───────────────────────────────────
heading2("Entity Schema Registry")

field_label("Description")
body(
    "A registry where operators declare the named entity types relevant to their agent "
    "deployment domain. Declarations include: entity type names and definitions (e.g. "
    "'patient identifier,' 'CUSIP,' 'procedure code'), acceptable transformation rules per "
    "entity type (e.g. capitalisation normalisation, accepted abbreviations, code expansions), "
    "and domain-specific severity mappings. Different industries require different entity "
    "vocabularies — the registry enables the evaluation system to apply domain-appropriate "
    "rules rather than general heuristics."
)

field_label("Enables")
bullet("REQ-03 Phase 2: Full Entity Accuracy evaluation — domain-specific entity type definitions and acceptable transformation rules")
bullet("REQ-10 Phase 2: Full Source Traceability evaluation — domain-specific entity enumeration and coverage completeness checks")
bullet("Eliminates unknown_not_computable (registry_not_configured) outcomes for REQ-03 and REQ-10 in regulated deployments")
bullet("Compliance-grade entity verification for healthcare (clinical terms, patient identifiers) and financial services (account numbers, security identifiers, regulatory references)")

field_label("Platform Capabilities Required")
bullet("Entity type declaration interface for operators with schema validation")
bullet("Entity schema storage per deployment with versioning and audit trail")
bullet("Entity extraction configuration applied at evaluation time per declared schema")
bullet("Transformation rule engine: acceptable abbreviations, code expansions, normalisation rules per entity type")
bullet("Integration with industry-standard entity vocabularies (ICD codes, LOINC, CUSIP, ISIN) for regulated deployments")

field_label("Evaluation Impact")
body(
    "Without this enhancement REQ-03 and REQ-10 operate using general entity categories "
    "derived from span data. With this enhancement the evaluator applies "
    "domain-specific entity definitions, transformation rules, and coverage completeness "
    "checks — the precision required for regulated industry compliance."
)

field_label("REQ-03 Phase 2 Requirement Extension")
body(
    "When the Entity Schema Registry is available, the Entity Accuracy check shall be "
    "extended as follows:"
)
bullet("The evaluator shall retrieve the declared entity schema for the deployment domain.")
bullet("The evaluator shall enumerate entities in the agent response according to the declared entity type definitions rather than general heuristics.")
bullet("For each entity, the evaluator shall apply the declared transformation rules for that entity type — acceptable abbreviations, code expansions, capitalisation normalisation — before determining whether the entity matches its source.")
bullet("Entity mismatches that fall within declared transformation rules shall be classified as no_hallucination. Entity mismatches that fall outside declared transformation rules shall retain their Phase 1 classification (minor or major).")
bullet("For regulated deployments with zero-tolerance entity types declared in the schema, any mismatch — regardless of degree — shall be classified as major_hallucination.")
bullet("The evaluator shall produce a completeness report: all entity types declared in the schema that were expected in the response but absent.")

field_label("REQ-10 Phase 2 Requirement Extension")
body(
    "When the Entity Schema Registry is available, the Source Traceability check shall be "
    "extended as follows:"
)
bullet("The evaluator shall enumerate all entities in the agent response according to the declared entity schema, using domain-specific entity type definitions to identify entities that general heuristics would miss.")
bullet("The evaluator shall apply declared transformation rules when tracing entities from response to source span — a transformation within declared rules is a valid match, not a traceability gap.")
bullet("The evaluator shall produce a coverage completeness report: entity types declared in the schema that were expected to appear in the response (based on the user query and tool outputs) but that are absent or untraceable.")
bullet("The source map produced by this check shall include entity type classification per the declared schema in addition to the span ID and span type.")
bullet("The evaluator shall no longer return unknown_not_computable (registry_not_configured) for domain-specific entity traceability once this enhancement is active.")

# ── Enhancement 4 — Regulated Industry Risk Profile Engine ───────────────────
heading2("Regulated Industry Risk Profile Engine")

field_label("Description")
body(
    "A configurable rule engine that applies regulatory risk profiles to the tiered evaluation "
    "strategy, automatically elevating specific checks from Tier 3 to mandatory Tier 2 or "
    "Tier 1 based on the agent deployment's declared industry and regulatory context. Profiles "
    "would be maintained and updated as regulatory requirements evolve."
)

field_label("Enables")
bullet("Automatic elevation of REQ-10 (Source Traceability) to mandatory Tier 2 for financial services and healthcare deployments")
bullet("Automatic elevation of REQ-09 (Confidence Calibration) to mandatory Tier 2 for deployments subject to suitability or clinical appropriateness requirements")
bullet("Automatic elevation of REQ-08 Phase 2 (Authorized Scope) to mandatory Tier 2 for regulated deployments")
bullet("Automatic elevation of REQ-03 Phase 2 (Entity Accuracy) severity thresholds to zero-tolerance for domain-critical entity types in regulated deployments")
bullet("Configurable severity thresholds per requirement per industry — e.g. REQ-04 Uncertainty Acknowledgment carries elevated severity in healthcare vs. non-regulated deployments")
bullet("Regulatory profile versioning as regulations evolve")

field_label("Platform Capabilities Required")
bullet("Industry and regulatory context declaration per agent deployment")
bullet("Profile rule engine with configurable tier elevation rules")
bullet("Profile management interface for operators and platform administrators")
bullet("Audit logging of which profile was applied to which evaluation")
bullet("Profile update workflow with change notification to deployment owners")

field_label("Evaluation Impact")
body(
    "Without this enhancement regulated deployments must manually treat Tier 3 as fully "
    "mandatory and configure severity thresholds outside the platform. With this enhancement "
    "the platform enforces regulatory requirements automatically and consistently across all "
    "agents in a regulated deployment."
)

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — REFERENCE IMPLEMENTATION APPENDIX
# ══════════════════════════════════════════════════════════════════════════════

heading1("Appendix — Reference Implementation: LG Travel Agent")

heading2("System Description")
body(
    "The LG Travel Agent is a LangGraph multi-agent application instrumented with Monocle "
    "v0.7.7 / OpenTelemetry. Architecture: one supervisor agent "
    "(okahu_demo_lg_agent_travel_supervisor) coordinating three sub-agents — air travel "
    "assistant (okahu_demo_lg_agent_air_travel_assistant), lodging assistant "
    "(okahu_demo_lg_agent_lodging_assistant), and a weather agent via MCP. Tools: "
    "okahu_demo_lg_tool_book_flight, okahu_demo_lg_tool_book_hotel, demo_get_weather. "
    "Model: gpt-4o via OpenAI API. Framework: LangGraph with langgraph_supervisor."
)

heading2("Trace Analysis Summary")
body(
    "Six production traces were analysed, all originally labelled no_hallucination by the "
    "prior evaluation template. Hallucinations were found in all six."
)

# 5-column trace table
trace_table = doc.add_table(rows=1, cols=5)
trace_table.alignment = WD_TABLE_ALIGNMENT.LEFT
trace_table.style = "Table Grid"

trace_headers = ["Trace ID", "User Query Summary", "Key Tool Output", "Agent Final Response", "Hallucination Present"]
hdr = trace_table.rows[0].cells
for i, h in enumerate(trace_headers):
    set_cell_bg(hdr[i], COLOR_HEADER)
    set_cell_borders(hdr[i], "1B2A4A", "6")
    p = hdr[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = COLOR_WHITE

trace_rows = [
    (
        "88748986",
        "Flight + hotel, Dallas to Paris TX",
        "Hotel tool: 'Hôtel République (Paris, France)'",
        "'Your hotel in Paris, Texas has been successfully booked.'",
        "Major — REQ-01, REQ-02, REQ-03, REQ-10",
    ),
    (
        "60a156dac",
        "Flight + hotel + weather, Atlanta to Paris TX",
        "Hotel tool: 'Hotel Republique (Paris, France)'",
        "'Reserved a hotel in Paris, Texas.'",
        "Major — REQ-01, REQ-02, REQ-03, REQ-10",
    ),
    (
        "182c5571",
        "Flight + hotel + weather, ATL to DFW, Paris TX",
        "Hotel tool: {}",
        "'Hotel in Paris, TX booked April 17–20.'",
        "Major — REQ-01, REQ-02, REQ-03, REQ-04, REQ-05",
    ),
    (
        "7263beb5",
        "Flight + hotel + weather, Atlanta to Nashville, Paris TN",
        "Hotel tool: {} — sub-agent exposed France",
        "Supervisor: 'Hotel de la Seine in Paris, Tennessee.'",
        "Major — REQ-01, REQ-02, REQ-03, REQ-06",
    ),
    (
        "da6d7367",
        "Flight + hotel + weather, Atlanta to Portland ME, Paris ME",
        "Hotel tool: {}",
        "'Hotel de la Seine in Paris, France.' (wrong country delivered)",
        "Major — REQ-01, REQ-02, REQ-03",
    ),
    (
        "df7648b4",
        "Flight + hotel + weather, ATL to DFW, Paris TX with dates",
        "Hotel tool: {}",
        "'Hotel de la Seine in Paris, France.' (wrong country delivered)",
        "Major — REQ-01, REQ-02, REQ-03",
    ),
]

for idx, (tid, query, tool_out, agent_resp, halluc) in enumerate(trace_rows):
    row = trace_table.add_row()
    cells = row.cells
    bg = COLOR_ROW1 if idx % 2 == 0 else COLOR_WHITE

    for ci, text in enumerate([tid, query, tool_out, agent_resp, halluc]):
        set_cell_bg(cells[ci], bg)
        set_cell_borders(cells[ci])
        cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cells[ci].paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(8.5)
        run.font.color.rgb = COLOR_BODY
        if ci == 0:
            run.bold = True
            run.font.color.rgb = COLOR_TEAL
        if ci == 4:
            run.font.color.rgb = RGBColor(0xA8, 0x1C, 0x1C)

for row in trace_table.rows:
    row.cells[0].width = Inches(0.75)
    row.cells[1].width = Inches(1.5)
    row.cells[2].width = Inches(1.6)
    row.cells[3].width = Inches(1.75)
    row.cells[4].width = Inches(0.8)

doc.add_paragraph()

heading2("Key Findings")
bullet("All six traces contained major hallucinations. None were detected by the prior evaluation template.")
bullet("The hotel booking tool (okahu_demo_lg_tool_book_hotel) returned Paris, France results regardless of the requested city — a systematic tool-level failure that the agent layer consistently failed to flag or correct.")
bullet("April 9 traces (88748986, 60a156dac): tool returned a named Paris, France hotel. Agents overrode the tool output and confirmed the user's requested city instead.")
bullet("April 10 traces (182c5571, 7263beb5, da6d7367, df7648b4): tool returned {}. Agents fabricated complete booking confirmations from empty tool responses. The supervisor system prompt was updated between trace sets to instruct exact relay of tool-returned hotel details — this caused the France result to surface rather than be covered up, but did not prevent the wrong booking.")
bullet("The system prompt evolution visible across trace dates confirms the development team was aware of the location mismatch problem and attempted to address it via prompt engineering. The fix surfaced the symptom without resolving the root cause.")
bullet("REQ-01 (Action Verification) and REQ-02 (Tool Output Faithfulness) together would have flagged all six traces under the requirements defined in this document.")
bullet("Weather tool was called with {city: 'Paris'} rather than {city: 'Paris, Texas'} across all traces where the user specified Texas — an observable scope drift detectable by REQ-08 Phase 1.")

doc.add_paragraph()


# ── Footer ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
add_paragraph_border_bottom(p, "C5CEDE")
p = doc.add_paragraph()
r = p.add_run(
    "Generated from analysis of Monocle/OpenTelemetry production traces for the LG Travel "
    "Agent (April 9–10, 2026). Requirements designed for the Okahu agentic observation "
    "platform. Phase 2 enhancements are planned capabilities, not current features."
)
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x7A, 0x8A, 0xA0)
r.italic = True


# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/Users/careyjames/Documents/GitHub/lg-travel-agent/Hallucination_Evaluation_Platform_Requirements.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
