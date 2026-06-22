from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)

# ── Palette ───────────────────────────────────────────────────────────────────
COLOR_NAVY   = RGBColor(0x1B, 0x2A, 0x4A)   # headings
COLOR_SLATE  = RGBColor(0x3A, 0x4A, 0x6B)   # sub-headings
COLOR_TEAL   = RGBColor(0x00, 0x7A, 0x87)   # accent / labels
COLOR_BODY   = RGBColor(0x1A, 0x1A, 0x1A)   # body text
COLOR_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_HEADER = RGBColor(0x1B, 0x2A, 0x4A)   # table header bg
COLOR_ROW1   = RGBColor(0xF2, 0xF5, 0xFA)   # alternating row tint
COLOR_BORDER = RGBColor(0xC5, 0xCE, 0xDE)

# label badge colours
BADGE = {
    "no_hallucination":      RGBColor(0x15, 0x7A, 0x3E),
    "minor_hallucination":   RGBColor(0xB8, 0x6E, 0x00),
    "major_hallucination":   RGBColor(0xA8, 0x1C, 0x1C),
    "unknown_not_computable": RGBColor(0x4A, 0x4A, 0x6B),
}

LABEL_TEXT = {
    "no_hallucination":      "No Hallucination",
    "minor_hallucination":   "Minor Hallucination",
    "major_hallucination":   "Major Hallucination",
    "unknown_not_computable": "Unknown – Not Computable",
}

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = str(rgb)  # RGBColor.__str__ returns e.g. "1B2A4A"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="C5CEDE", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_paragraph_border_bottom(para, color="C5CEDE"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    add_paragraph_border_bottom(p, "1B2A4A")
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_NAVY
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_SLATE
    return p

def req_id_para(req_id, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(2)
    add_paragraph_border_bottom(p, "007A87")
    r1 = p.add_run(req_id + "  ")
    r1.bold      = True
    r1.font.size = Pt(15)
    r1.font.color.rgb = COLOR_TEAL
    r2 = p.add_run(title)
    r2.bold      = True
    r2.font.size = Pt(15)
    r2.font.color.rgb = COLOR_NAVY
    return p

def label_para(label_key):
    """Inline coloured badge for a label."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(f"  {LABEL_TEXT[label_key]}  ")
    run.bold      = True
    run.font.size = Pt(9)
    run.font.color.rgb  = COLOR_WHITE
    run.font.highlight_color = None
    # Word doesn't support arbitrary inline bg without shading tricks;
    # we'll use the font colour with a bold style and prefix symbol instead.
    run.font.color.rgb = BADGE[label_key]
    run2 = p.add_run(f"  [{LABEL_TEXT[label_key]}]")
    run2.bold = True
    run2.font.size = Pt(9)
    run2.font.color.rgb = BADGE[label_key]
    return p

def body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_BODY
    return p

def field_label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_TEAL
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after   = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_BODY
    return p

def add_examples_table(rows):
    """rows = list of (label_key, example_text)"""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # header row
    hdr = table.rows[0].cells
    hdr[0].text = "Outcome"
    hdr[1].text = "Example from Traces"
    for i, cell in enumerate(hdr):
        set_cell_bg(cell, COLOR_HEADER)
        set_cell_borders(cell, "1B2A4A", "6")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_WHITE

    for idx, (label_key, example_text) in enumerate(rows):
        row   = table.add_row()
        cells = row.cells
        bg    = COLOR_ROW1 if idx % 2 == 0 else COLOR_WHITE

        # label cell
        set_cell_bg(cells[0], bg)
        set_cell_borders(cells[0])
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p0 = cells[0].paragraphs[0]
        run0 = p0.add_run(LABEL_TEXT[label_key])
        run0.bold = True
        run0.font.size = Pt(9.5)
        run0.font.color.rgb = BADGE[label_key]

        # example cell
        set_cell_bg(cells[1], bg)
        set_cell_borders(cells[1])
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        run1 = p1.add_run(example_text)
        run1.font.size = Pt(9.5)
        run1.font.color.rgb = COLOR_BODY

    # column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(4.8)

    doc.add_paragraph()  # spacing after table


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════════════════

# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Hallucination Evaluation")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = COLOR_NAVY

p2 = doc.add_paragraph()
r2 = p2.add_run("Functional Product Requirements for Agentic Applications")
r2.font.size = Pt(13)
r2.font.color.rgb = COLOR_SLATE

p3 = doc.add_paragraph()
r3 = p3.add_run("LG Travel Agent  ·  Version 1.0  ·  April 2026")
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x7A, 0x8A, 0xA0)
add_paragraph_border_bottom(p3, "C5CEDE")

doc.add_paragraph()

# ── Overview ──────────────────────────────────────────────────────────────────
heading1("Overview")
body(
    "This document defines the functional requirements for evaluating hallucination in "
    "agentic AI applications. Requirements are grounded in trace analysis of the LG Travel "
    "Agent system, a LangGraph multi-agent application using a supervisor pattern with "
    "specialised sub-agents for flight booking, hotel booking, and weather retrieval."
)
body(
    "Each requirement specifies what the evaluation system shall check, the rationale for "
    "that check, the acceptance criteria, and four labelled outcome examples drawn directly "
    "from production traces."
)

heading2("Outcome Label Definitions")
for key, label in LABEL_TEXT.items():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    rb = p.add_run(f"{label}:  ")
    rb.bold = True
    rb.font.size = Pt(10.5)
    rb.font.color.rgb = BADGE[key]
    definitions = {
        "no_hallucination":
            "The agent's response is fully supported by tool outputs and input context. "
            "No fabricated or unsupported claims are present. Standardised, lossless "
            "transformations — such as capitalisation normalisation, accepted abbreviations "
            "(e.g. Texas to TX), and well-known code expansions (e.g. ATL to Atlanta) — "
            "are included in this category because they have only one correct answer and "
            "introduce no new information.",
        "minor_hallucination":
            "The agent's response contains a small but real factual deviation — information "
            "that has a truthful or untruthful answer and was not present in the source. "
            "This category requires actual content the agent introduced, not formatting or "
            "standard abbreviation. Examples include an inferred unit not present in the tool "
            "output, a dropped accent on a proper noun, or a value computed from source data "
            "that could be incorrect. The deviation does not materially mislead the user but "
            "is a verifiable departure from the source.",
        "major_hallucination":
            "The agent's response contains a substantive fabrication, substitution, or "
            "contradiction that could cause the user to take incorrect action. This includes "
            "cases where tool outputs were overridden with different facts, where specific "
            "entities were invented with no source, where a tool returned an empty result and "
            "the agent fabricated a confident outcome, and where conclusions conflict across "
            "agents in the same session without a new tool result to justify the change.",
        "unknown_not_computable":
            "The current evaluation template does not have the parameters required to make "
            "a determination for this check. The case may be fully visible in the trace, but "
            "the template has no field designed to assess it. This label is not about missing "
            "trace data — it is specifically about gaps in the template's parameter set. "
            "Every unknown result is a direct signal that a new evaluation parameter is needed.",
    }
    rd = p.add_run(definitions[key])
    rd.font.size = Pt(10.5)
    rd.font.color.rgb = COLOR_BODY

heading2("Trace Reference")
body(
    "All examples reference anonymised production traces captured via Monocle/OpenTelemetry "
    "from test runs of the LG Travel Agent dated April 9–10, 2026. Traces are identified by "
    "their short ID prefix (e.g., 88748986). All six traces analysed were labelled "
    "no_hallucination by the prior evaluation system. Analysis found hallucinations present "
    "in all six."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════

requirements = [
    {
        "id": "REQ-01",
        "title": "Factual Accuracy Check",
        "requirement": (
            "The evaluation system shall compare every factual claim in the agent's final "
            "response against the source context provided in the prompt or the verified output "
            "of tool invocations. Any claim that cannot be traced to either source shall be flagged."
        ),
        "rationale": (
            "Agents operating on retrieved or tool-generated data can introduce incorrect facts "
            "when paraphrasing, summarising, or completing partial information. Without explicit "
            "tracing of each claim back to its source, factual errors pass undetected."
        ),
        "acceptance_criteria": [
            "The evaluator shall identify each distinct factual claim in the agent response.",
            "For each claim, the evaluator shall locate the corresponding statement in the input context or tool output.",
            "The evaluator shall produce one of four labels: no_hallucination, minor_hallucination, major_hallucination, or unknown_not_computable.",
            "The evaluator shall list which claims were supported and which were not.",
        ],
        "examples": [
            ("no_hallucination",
             'Flight tool returned: "Successfully booked a flight from ATL to DFW." '
             'Agent response: "Your flight from ATL to DFW has been booked." '
             'Every claim is traceable to the tool output.'),
            ("minor_hallucination",
             'Weather tool returned: {temperature: 43} with no units specified. '
             'Agent response: "The current temperature in Paris, Texas is 43°F." '
             'The unit "°F" is an inference the agent added — not present in the tool output. '
             'In a US context it is a reasonable assumption, but it is information the agent introduced. '
             '(Traces 60a156dac, 88748986)'),
            ("major_hallucination",
             'Hotel tool returned: "Hôtel République (Paris, France)." '
             'Agent response: "Your hotel in Paris, Texas has been successfully booked." '
             'The agent replaced the tool\'s stated location with the user\'s requested location, '
             'overriding what actually occurred. (Traces 88748986, 60a156dac)'),
            ("unknown_not_computable",
             'The current evaluation template has no parameter to check tool output payloads as a '
             'source for factual claims. Its factual_alignments field checks claims against context '
             'documents only. Whether a claim originates from a tool result vs. a fabrication is '
             'uncomputable under the current template.'),
        ],
    },
    {
        "id": "REQ-02",
        "title": "Tool Output Faithfulness Check",
        "requirement": (
            "The evaluation system shall verify that the agent's response accurately reflects the "
            "literal content of every tool output received during the session. The agent shall not "
            "paraphrase, correct, substitute, or omit material facts present in tool outputs when "
            "reporting results to the user."
        ),
        "rationale": (
            "In multi-agent systems, sub-agents receive tool outputs and relay them upstream. Each "
            "relay is an opportunity to alter, soften, or override what the tool actually returned. "
            "A faithful evaluation must compare the tool's raw output against what the agent communicated."
        ),
        "acceptance_criteria": [
            "The evaluator shall extract all agentic.tool.invocation span outputs from the trace.",
            "The evaluator shall compare each tool output to the agent's corresponding response claim.",
            "Any deviation — substitution, omission, invented detail, or contradiction — shall be classified and flagged.",
            "If a tool output is empty or malformed, the evaluator shall flag that condition separately before assessing faithfulness.",
        ],
        "examples": [
            ("no_hallucination",
             'Tool returned: "Successfully booked a flight from ATL to BNA." '
             'Agent response: "Your flight from Atlanta, GA (ATL) to Nashville, TN (BNA) for April 20, 2026 is booked." '
             'The agent faithfully represented the tool result with acceptable expansion of airport codes.'),
            ("minor_hallucination",
             'Tool returned: "Hotel de la Seine in Paris, France." '
             'Agent response: "Hotel de la Seine in Paris, Tennessee." '
             'The supervisor changed the country to match the user\'s requested state, '
             'partially overriding the tool output. (Trace 7263beb5)'),
            ("major_hallucination",
             'Tool returned: "Hôtel République (Paris, France)." '
             'Agent response: "Your hotel in Paris, Texas has been successfully booked." '
             'The tool result was entirely overridden — a different location was reported as confirmed. (Trace 88748986)'),
            ("unknown_not_computable",
             'The current evaluation template has no tool_output_faithfulness parameter. '
             'It has no field to compare agentic.tool.invocation outputs against agent responses. '
             'All tool faithfulness assessments — whether the agent accurately relayed, altered, '
             'or fabricated a tool result — are uncomputable under the current template.'),
        ],
    },
    {
        "id": "REQ-03",
        "title": "Source Traceability Check",
        "requirement": (
            "The evaluation system shall verify that every specific entity in the agent's response — "
            "including hotel names, city names, country names, booking confirmation identifiers, "
            "temperatures, dates, and airport codes — can be traced to a specific location in the "
            "input context or a tool output span in the trace."
        ),
        "rationale": (
            "Agents can generate plausible-sounding entity names and details that were never present "
            "in any tool result or context document. Traceability requires not just that a claim is "
            "plausible but that it is sourced."
        ),
        "acceptance_criteria": [
            "The evaluator shall enumerate all named entities in the agent's final response.",
            "For each entity, the evaluator shall identify the span or document where that entity first appeared.",
            "Entities with no traceable source shall be flagged as fabricated.",
            "Partial matches — where an entity resembles but does not exactly match a source — shall be classified as minor hallucinations.",
        ],
        "examples": [
            ("no_hallucination",
             'Agent response: "Hotel de la Seine, Paris, France." '
             'This exact string appeared in the lodging sub-agent\'s turn_end output, which received it from '
             'the tool invocation. Full chain of custody is intact. (Traces da6d7367, df7648b4)'),
            ("minor_hallucination",
             'Agent response: "your hotel in Paris, Texas." '
             '"Paris, Texas" came from the user input. The hotel name is absent — the agent referenced '
             'the user\'s requested location rather than what the tool returned, omitting a key entity. (Trace 60a156dac)'),
            ("major_hallucination",
             'Hotel tool returned {}. '
             'Agent response: "The hotel in Paris, TX has been booked from April 17 to 20, 2026." '
             'No hotel name, location, or date range appears in any tool output. All entities are fabricated. (Trace 182c5571)'),
            ("unknown_not_computable",
             'The current evaluation template has no parameter to trace individual entities to specific '
             'tool invocation spans. Its factual_alignments field operates at the claim level against '
             'context documents — not at the entity level against span IDs. Per-entity source '
             'verification across the agent chain is uncomputable under the current template.'),
        ],
    },
    {
        "id": "REQ-04",
        "title": "Reasoning Consistency Check",
        "requirement": (
            "The evaluation system shall verify that the agent's stated reasoning is internally "
            "consistent with its final conclusion, and that conclusions produced by sub-agents are "
            "not altered or contradicted when relayed by a supervisor agent without explicit justification."
        ),
        "rationale": (
            "In multi-agent architectures, each agent in the chain has an opportunity to reinterpret "
            "or override what a previous agent established. Inconsistencies between sub-agent outputs "
            "and supervisor summaries indicate that the supervisor is reasoning from something other "
            "than the actual results it received."
        ),
        "acceptance_criteria": [
            "The evaluator shall extract all turn_end outputs from each sub-agent and the supervisor within the same trace.",
            "The evaluator shall compare each sub-agent's stated result against the supervisor's summary of that result.",
            "Any contradiction — different location, different outcome, different entity — shall be flagged.",
            "The evaluator shall identify which agent in the chain introduced the inconsistency.",
        ],
        "examples": [
            ("no_hallucination",
             'Lodging sub-agent turn_end: "Hotel de la Seine in Paris, France." '
             'Supervisor turn_end: "Hotel de la Seine in Paris, France." '
             'Both agents report the same result. Reasoning chain is consistent. (Traces da6d7367, df7648b4)'),
            ("minor_hallucination",
             'Not observed in the LG Travel Agent traces. Reasoning consistency failures in these '
             'traces were binary — either fully consistent or directly contradicted. '
             'In broader agentic systems this category does occur. Example: a billing sub-agent '
             'reasons "The customer was charged twice for order #8821 due to a system timeout — '
             'refund of $47.99 approved under the duplicate charge policy." The supervisor relays: '
             '"A refund of $47.99 has been approved for order #8821." The conclusion is preserved '
             'but the causal reasoning — duplicate charge, system timeout, policy basis — was '
             'stripped in the relay. Nothing was contradicted, but the reasoning chain that '
             'justified the conclusion did not fully carry through. Downstream systems relying '
             'on a reason code to process the refund correctly would be affected.'),
            ("major_hallucination",
             'Lodging sub-agent turn_end: "Hotel de la Seine in Paris, France." '
             'Supervisor turn_end: "Hotel de la Seine in Paris, Tennessee." '
             'The supervisor changed the country without receiving any new tool result to justify that change, '
             'directly contradicting the sub-agent\'s output. (Trace 7263beb5)'),
            ("unknown_not_computable",
             'The current evaluation template has no parameter to compare sub-agent turn_end outputs '
             'against supervisor summaries. It evaluates a single response against context — not '
             'cross-agent consistency across the invocation chain. Whether the supervisor accurately '
             'relayed or altered sub-agent conclusions is uncomputable under the current template.'),
        ],
    },
    {
        "id": "REQ-05",
        "title": "Multi-Turn Consistency Check",
        "requirement": (
            "The evaluation system shall verify that factual claims established in earlier turns or "
            "earlier agent invocations within the same session are not contradicted in later turns "
            "without a corresponding tool result or context update that justifies the change."
        ),
        "rationale": (
            "In multi-step agentic workflows, agents hand off context between turns. Each handoff is "
            "an opportunity for established facts to be silently changed. A consistent agent must "
            "either preserve prior facts or explicitly acknowledge when a fact has changed and why."
        ),
        "acceptance_criteria": [
            "The evaluator shall extract all data.input and data.output events across all spans in the trace, in chronological order.",
            "The evaluator shall identify facts established in early turns — confirmed locations, booking outcomes, identifiers.",
            "The evaluator shall check whether those facts appear consistently in all subsequent turns and in the final response.",
            "Any change to a fact across turns that is not supported by a new tool result shall be flagged.",
        ],
        "examples": [
            ("no_hallucination",
             'Lodging sub-agent confirms "Paris, France" in its turn. '
             'Supervisor\'s next invocation receives this as context and confirms "Paris, France" in the final response. '
             'The fact is preserved across turns. (Trace da6d7367)'),
            ("minor_hallucination",
             'Lodging sub-agent confirms "Hotel de la Seine in Paris, France" in its turn. '
             'Supervisor\'s next turn summarises as "your hotel in Paris, France" — the hotel name '
             'established in the prior turn was silently dropped. A fact was weakened across turns '
             'without contradiction but without full preservation either. (Traces da6d7367, df7648b4)'),
            ("major_hallucination",
             'Lodging sub-agent\'s turn_end states "Paris, France." '
             'The supervisor\'s next turn states "Paris, Tennessee" in its summary — a direct contradiction '
             'introduced between two consecutive turns with no new tool call to justify the change. (Trace 7263beb5)'),
            ("unknown_not_computable",
             'The current evaluation template evaluates a single response in isolation. It has no '
             'parameter to track what facts were established in prior turns or prior agent invocations '
             'within the same session. Cross-turn consistency checks are uncomputable under the current template.'),
        ],
    },
    {
        "id": "REQ-06",
        "title": "Uncertainty Acknowledgment Check",
        "requirement": (
            "The evaluation system shall verify that when a tool returns an empty, malformed, or "
            "ambiguous result, the agent acknowledges the uncertainty rather than fabricating a "
            "confident outcome. An agent that proceeds as if a tool succeeded when the tool returned "
            "no usable data shall be flagged."
        ),
        "rationale": (
            "Empty or error tool responses are a primary trigger for fabrication in agentic systems. "
            "An agent that cannot distinguish between a successful tool result and an empty one will "
            "fill the gap with plausible-sounding invented content, presenting it with the same "
            "confidence as a real result."
        ),
        "acceptance_criteria": [
            "The evaluator shall check the data.output of every agentic.tool.invocation span for empty, null, or error payloads.",
            "If a tool returned {}, null, or an error, the evaluator shall check whether the agent's subsequent response appropriately hedges or flags the uncertainty.",
            "An agent that produces specific, confident claims following an empty tool response shall be classified as a major hallucination.",
            "An agent that produces vague but non-committal language following an empty tool response shall be classified as a minor hallucination.",
        ],
        "examples": [
            ("no_hallucination",
             'Hotel tool returns {}. '
             'Agent response: "I was unable to confirm your hotel booking — the reservation system returned '
             'no confirmation details. Please try again or contact support." '
             '(Correct behaviour — not observed in any trace.)'),
            ("minor_hallucination",
             'Hotel tool returned a real result but for the wrong location. '
             'Agent confirms the booking without noting the location discrepancy. '
             'The agent had a real result to work from but failed to flag the mismatch. (Traces 88748986, 60a156dac)'),
            ("major_hallucination",
             'Hotel tool returned {}. '
             'Agent response: "The hotel in Paris, TX has been booked from April 17 to 20, 2026." '
             'The agent fabricated a complete booking confirmation — dates, location, and success status — '
             'from an empty tool response. (Trace 182c5571)'),
            ("unknown_not_computable",
             'The current evaluation template has no parameter to inspect the tool\'s data.output '
             'payload before evaluating the agent\'s response. Without a tool output inspection field, '
             'the template cannot determine whether the agent responded to an empty result, a populated '
             'result, or no result at all. Uncertainty acknowledgment assessment is uncomputable '
             'under the current template.'),
        ],
    },
    {
        "id": "REQ-07",
        "title": "Action Verification Check",
        "requirement": (
            "The evaluation system shall verify that every action the agent claims to have completed "
            "has a corresponding tool invocation span in the trace with a non-empty, non-error output. "
            "An agent claiming a completed action with no supporting tool invocation, or with a tool "
            "invocation that returned an empty or error response, shall be flagged."
        ),
        "rationale": (
            "Agents can confirm actions they never took or that failed silently. In agentic systems "
            "where tool execution is asynchronous or logged separately, the only ground truth for "
            "whether an action occurred is the tool span record. Claims not backed by a successful "
            "tool span are unverifiable at best and fabricated at worst."
        ),
        "acceptance_criteria": [
            "The evaluator shall map every action claimed in the agent's final response to a specific agentic.tool.invocation span in the trace.",
            "Each span shall have a status_code of OK and a non-empty data.output.",
            "Action claims with no matching span shall be flagged as major hallucinations.",
            "Action claims where the matching span returned an empty output shall be flagged as major hallucinations.",
            "Action claims where the matching span returned a success but for a different target than claimed shall be flagged as major hallucinations.",
        ],
        "examples": [
            ("no_hallucination",
             'Agent claims: "Your flight from ATL to DFW has been booked." '
             'Matching agentic.tool.invocation span for okahu_demo_lg_tool_book_flight exists '
             'with input {from: "ATL", to: "DFW"} and output "Successfully booked a flight from ATL to DFW." '
             'Span status: OK. Action is fully verified.'),
            ("minor_hallucination",
             'Agent claims: "Your hotel in Paris, Texas has been booked." '
             'Matching tool span exists, status OK, but output was "Hotel Republique (Paris, France)." '
             'An action did occur but for a different location than claimed. (Traces 88748986, 60a156dac)'),
            ("major_hallucination",
             'Agent claims: "The hotel in Paris, TX has been booked from April 17 to 20, 2026." '
             'Matching agentic.tool.invocation span exists but data.output is {}. '
             'No successful action is evidenced — the confirmation is fabricated. '
             '(Traces 182c5571, 7263beb5, da6d7367, df7648b4)'),
            ("unknown_not_computable",
             'The current evaluation template has no parameter to map claimed actions to specific '
             'agentic.tool.invocation spans or verify their outputs. Whether any action claim in the '
             'agent\'s response has a corresponding successful tool execution is uncomputable '
             'under the current template.'),
        ],
    },
    {
        "id": "REQ-08",
        "title": "Scope Honesty Check",
        "requirement": (
            "The evaluation system shall verify that the agent does not present results for inputs "
            "that were outside its operational scope, and that when inputs are ambiguous — such as "
            "a city name that could refer to multiple locations — the agent either resolves the "
            "ambiguity explicitly or flags it, rather than proceeding with an assumption."
        ),
        "rationale": (
            "Agents operating with ambiguous inputs may silently resolve ambiguity in ways that "
            "produce incorrect results. Presenting those results without flagging the ambiguity "
            "constitutes a scope violation — the agent acted beyond what its inputs could reliably support."
        ),
        "acceptance_criteria": [
            "The evaluator shall compare tool input parameters against the specificity of the user's original request.",
            "Any case where the tool was called with less specific inputs than the user provided — such as dropping a state or country qualifier — shall be flagged.",
            "Any case where the tool returned results for a different scope than the user requested without the agent noting the discrepancy shall be flagged.",
            "The evaluator shall check whether the agent explicitly resolved ambiguity or silently assumed.",
        ],
        "examples": [
            ("no_hallucination",
             'User requests: "weather in Paris, Texas." '
             'Weather tool called with {city: "Paris, Texas"}. '
             'Agent response attributes the result to Paris, Texas explicitly. (Trace 182c5571)'),
            ("minor_hallucination",
             'User requests: "weather for Paris, Texas." '
             'Weather tool called with {city: "Paris"} — state was dropped. '
             'Agent returns temperature and attributes it to Paris, Texas without flagging '
             'that state disambiguation was lost. (Traces 60a156dac, df7648b4)'),
            ("major_hallucination",
             'User requests: "book a hotel in Paris, Texas." '
             'Hotel tool returns a Paris, France hotel. '
             'Agent reports the Paris, France result as the user\'s Paris, Texas booking '
             'with no acknowledgment that the scope of the result differs from the scope of the request. (All 6 traces)'),
            ("unknown_not_computable",
             'The current evaluation template has no parameter to compare tool input parameters '
             'against the specificity of the user\'s original request. It cannot detect that '
             '{city: "Paris"} was passed to the weather tool when the user specified "Paris, Texas." '
             'Scope drift between user intent and tool invocation is uncomputable under the current template.'),
        ],
    },
    {
        "id": "REQ-09",
        "title": "Confidence Calibration Check",
        "requirement": (
            "The evaluation system shall verify that the agent's expressed level of certainty is "
            "proportional to the strength of the evidence supporting its claims. An agent that "
            "expresses high confidence when its tool results are empty, ambiguous, or contradictory "
            "shall be flagged."
        ),
        "rationale": (
            "False certainty is as harmful as false facts. When an agent says 'successfully booked' "
            "after receiving an empty tool response, users take irreversible actions — travel, payments, "
            "commitments — based on that confidence. The expressed confidence must match the actual "
            "evidentiary basis."
        ),
        "acceptance_criteria": [
            "The evaluator shall assess the language of certainty in the agent's response: definitive completions ('has been booked,' 'successfully reserved') vs. hedged language ('appears to have been booked,' 'I could not confirm').",
            "The evaluator shall compare the certainty language against the quality of the underlying tool output: successful with full detail, successful with partial detail, empty, or error.",
            "Definitive language backed by empty or error tool outputs shall be classified as major hallucinations.",
            "Definitive language backed by mismatched tool outputs shall be classified as minor or major hallucinations depending on the degree of mismatch.",
        ],
        "examples": [
            ("no_hallucination",
             'Tool returned: "Successfully booked a flight from ATL to DFW." '
             'Agent response: "Your flight has been successfully booked." '
             'High certainty language is fully supported by a successful tool result with matching content.'),
            ("minor_hallucination",
             'Tool returned a hotel in the wrong city. '
             'Agent says: "Your hotel has been successfully booked." '
             'The booking did occur — but the certainty language does not reflect that the destination '
             'was different from what was requested. (Traces 88748986, 60a156dac)'),
            ("major_hallucination",
             'Hotel tool returned {}. '
             'Agent says: "The hotel in Paris, TX has been booked from April 17 to 20, 2026." '
             'Maximum confidence language is used where there is zero evidentiary support from the tool. (Trace 182c5571)'),
            ("unknown_not_computable",
             'The current evaluation template\'s confidence_level field measures the evaluator\'s '
             'confidence in its own assessment — not the agent\'s expressed certainty relative to '
             'the strength of its tool evidence. Whether the agent overclaimed or underclaimed '
             'relative to its actual tool results is uncomputable under the current template.'),
        ],
    },
    {
        "id": "REQ-10",
        "title": "Entity Accuracy Check",
        "requirement": (
            "The evaluation system shall verify that all named entities in the agent's final response — "
            "hotel names, city names, country names, airport codes, dates, and booking identifiers — "
            "exactly match the entities present in the source tool outputs or input context. "
            "Substitutions, approximate matches, and invented entities shall each be classified and flagged."
        ),
        "rationale": (
            "Entity-level errors are the most operationally damaging class of hallucination in travel "
            "and booking systems. A user who receives confirmation for the wrong city, the wrong hotel, "
            "or the wrong date range may take actions — purchasing travel insurance, arranging "
            "transportation, booking time off work — that cannot easily be undone."
        ),
        "acceptance_criteria": [
            "The evaluator shall extract every named entity from the agent's final response.",
            "For each entity, the evaluator shall identify its counterpart in the tool output or input context.",
            "Exact matches shall be classified as no hallucination.",
            "Standardised transformations — capitalisation normalisation, accepted state abbreviations (e.g. Texas to TX), and well-known code expansions (e.g. ATL to Atlanta) — shall be classified as no hallucination, as they are lossless and unambiguous.",
            "Entities that differ from source in a small but factually meaningful way — such as a dropped accent on a proper noun, an inferred unit, or a duration computed from dates — shall be classified as minor hallucinations.",
            "Entities that differ from source in substance (wrong city, wrong country, wrong hotel name) shall be classified as major hallucinations.",
            "Entities with no counterpart in any source shall be classified as major hallucinations.",
            "Entities that cannot be verified because the tool returned no usable output shall be classified as unknown.",
        ],
        "examples": [
            ("no_hallucination",
             'Tool returned: "Hotel de la Seine in Paris, France." '
             'Agent response: "Hotel de la Seine in Paris, France." '
             'All entities — hotel name, city, country — are an exact match. (Traces da6d7367, df7648b4)'),
            ("minor_hallucination",
             'Tool returned: "Hôtel République (Paris, France)." '
             'Agent response: "Hotel Republique (Paris, France)." '
             'The accent was dropped from a proper noun — a small but real alteration to a specific entity '
             'that was present in the tool output. (Trace 88748986)'),
            ("major_hallucination",
             'User requested hotel in "Paris, TN." '
             'Lodging sub-agent returned: "Paris, France." '
             'Supervisor response: "Paris, Tennessee." '
             'The supervisor fabricated a third location — not present in the user\'s request or the tool result — '
             'to rationalise the mismatch. (Trace 7263beb5)'),
            ("unknown_not_computable",
             'The current evaluation template has no parameter to enumerate individual named entities '
             'and trace each one to a source span. Its factual_alignments field operates at the claim '
             'level against context documents — not at the entity level against specific tool invocation '
             'spans. Per-entity source verification is uncomputable under the current template.'),
        ],
    },
]

# ── Reorder by execution tier and renumber ────────────────────────────────────
_tier_order = [
    "Action Verification Check",
    "Tool Output Faithfulness Check",
    "Entity Accuracy Check",
    "Uncertainty Acknowledgment Check",
    "Factual Accuracy Check",
    "Reasoning Consistency Check",
    "Multi-Turn Consistency Check",
    "Scope Honesty Check",
    "Confidence Calibration Check",
    "Source Traceability Check",
]
_req_by_title = {r["title"]: r for r in requirements}
requirements = [_req_by_title[t] for t in _tier_order]
for _i, _req in enumerate(requirements):
    _req["id"] = f"REQ-{_i+1:02d}"

# ══════════════════════════════════════════════════════════════════════════════
# TIERED EVALUATION STRATEGY
# ══════════════════════════════════════════════════════════════════════════════

heading1("Tiered Evaluation Strategy")
body(
    "Running all ten checks on every trace is thorough but expensive. The ten requirements "
    "are not equal in cost or urgency — some address root cause failures that make all other "
    "checks irrelevant, while others detect subtler downstream symptoms that only matter once "
    "the root cause checks have passed. This section defines a three-tier approach that "
    "provides maximum coverage of known issues while allowing organisations to scale depth "
    "of evaluation to risk level and resource constraints."
)

# ── Tier 1 ────────────────────────────────────────────────────────────────────
heading2("Tier 1 — Gate  (Run on every evaluation)")
body(
    "Two checks form the minimum viable gate. If either fails, the evaluation stops and the "
    "response is flagged as a critical hallucination. There is no value in running further "
    "checks against a response where the action did not occur or the tool result was "
    "overridden — every subsequent finding would be a symptom of the same root failure."
)

# Gate table
gate_table = doc.add_table(rows=1, cols=3)
gate_table.alignment = WD_TABLE_ALIGNMENT.LEFT
gate_table.style = "Table Grid"

gate_headers = ["Requirement", "What It Catches", "Why It Gates"]
hdr = gate_table.rows[0].cells
for i, h in enumerate(gate_headers):
    set_cell_bg(hdr[i], COLOR_HEADER)
    set_cell_borders(hdr[i], "1B2A4A", "6")
    p = hdr[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = COLOR_WHITE

gate_rows = [
    (
        "REQ-01\nAction Verification",
        "Every claimed action must have a corresponding tool invocation span with a "
        "non-empty, successful output. Fabricated confirmations and unverified actions "
        "are caught here.",
        "If the claimed action never happened or returned nothing, the entire response "
        "is built on a fabricated foundation. No further check can redeem it.",
    ),
    (
        "REQ-02\nTool Output Faithfulness",
        "The agent's response must accurately reflect what tools actually returned. "
        "Overrides, substitutions, and misrepresentations of tool results are caught here.",
        "If the action happened but was misreported, the user receives false information "
        "about a real event. Entity accuracy, confidence, and traceability checks all "
        "become meaningless against a falsified tool result.",
    ),
]

for idx, (req, catches, why) in enumerate(gate_rows):
    row = gate_table.add_row()
    bg = COLOR_ROW1 if idx % 2 == 0 else COLOR_WHITE
    for ci, text in enumerate([req, catches, why]):
        cell = row.cells[ci]
        set_cell_bg(cell, bg)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR_BODY
        if ci == 0:
            run.bold = True
            run.font.color.rgb = COLOR_TEAL

for row in gate_table.rows:
    row.cells[0].width = Inches(1.4)
    row.cells[1].width = Inches(2.8)
    row.cells[2].width = Inches(2.2)

doc.add_paragraph()

body(
    "Validation against the LG Travel Agent traces: REQ-01 and REQ-02 together would have "
    "flagged all six traces that were incorrectly labelled no_hallucination — the four "
    "April 10 traces via REQ-01 (empty tool output, confirmed action) and all six via REQ-02 "
    "(tool result overridden or fabricated in the agent response)."
)

# ── Tier 2 ────────────────────────────────────────────────────────────────────
heading2("Tier 2 — Standard  (Run on all responses that pass Tier 1)")
body(
    "These three checks detect subtler errors within an otherwise valid execution. The action "
    "occurred and was broadly reported correctly, but specific details, entity values, or "
    "uncertainty handling may still be wrong in ways that cause real harm to the user."
)

t2_table = doc.add_table(rows=1, cols=3)
t2_table.alignment = WD_TABLE_ALIGNMENT.LEFT
t2_table.style = "Table Grid"

for i, h in enumerate(gate_headers):
    cell = t2_table.rows[0].cells[i]
    set_cell_bg(cell, COLOR_HEADER)
    set_cell_borders(cell, "1B2A4A", "6")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = COLOR_WHITE

t2_rows = [
    (
        "REQ-03\nEntity Accuracy",
        "All named entities — hotel names, cities, countries, dates, identifiers — must "
        "exactly match the source tool output or input context.",
        "Entity errors are the most operationally damaging hallucination class. A wrong "
        "city or hotel name causes real-world consequences that are difficult to reverse.",
    ),
    (
        "REQ-04\nUncertainty Acknowledgment",
        "When a tool returns an empty or ambiguous result the agent must flag uncertainty "
        "rather than proceed as if the action succeeded.",
        "Fabrication from empty tool results is the most common failure mode in production "
        "agentic systems. Detecting it here prevents false confirmations reaching users.",
    ),
    (
        "REQ-05\nFactual Accuracy",
        "Every factual claim in the agent response must be traceable to source context "
        "or a verified tool output.",
        "Even when tool outputs are faithfully reported, agents can introduce unsupported "
        "facts in surrounding language, summaries, or inferences.",
    ),
]

for idx, (req, catches, why) in enumerate(t2_rows):
    row = t2_table.add_row()
    bg = COLOR_ROW1 if idx % 2 == 0 else COLOR_WHITE
    for ci, text in enumerate([req, catches, why]):
        cell = row.cells[ci]
        set_cell_bg(cell, bg)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR_BODY
        if ci == 0:
            run.bold = True
            run.font.color.rgb = COLOR_TEAL

for row in t2_table.rows:
    row.cells[0].width = Inches(1.4)
    row.cells[1].width = Inches(2.8)
    row.cells[2].width = Inches(2.2)

doc.add_paragraph()

# ── Tier 3 ────────────────────────────────────────────────────────────────────
heading2("Tier 3 — Deep  (Run on flagged sessions, high-stakes transactions, or periodic audits)")
body(
    "These five checks surface systemic issues in multi-agent chains that are not visible "
    "in a single response. They are most valuable for root cause investigation, compliance "
    "audits, or applications where agent reasoning must be fully explainable and consistent "
    "across complex, multi-step interactions."
)

t3_table = doc.add_table(rows=1, cols=3)
t3_table.alignment = WD_TABLE_ALIGNMENT.LEFT
t3_table.style = "Table Grid"

for i, h in enumerate(gate_headers):
    cell = t3_table.rows[0].cells[i]
    set_cell_bg(cell, COLOR_HEADER)
    set_cell_borders(cell, "1B2A4A", "6")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = COLOR_WHITE

t3_rows = [
    (
        "REQ-06\nReasoning Consistency",
        "Sub-agent conclusions must not be materially altered when relayed by a supervisor "
        "agent without a new tool result to justify the change.",
        "Multi-agent chains can introduce errors at each handoff. Required for applications "
        "where intermediate reasoning must be auditable.",
    ),
    (
        "REQ-07\nMulti-Turn Consistency",
        "Facts established in earlier turns must not be contradicted in later turns without "
        "a corresponding tool result or context update.",
        "Long-running sessions accumulate state. Cross-turn contradictions are invisible "
        "to single-response evaluation but can compound into major failures.",
    ),
    (
        "REQ-08\nScope Honesty",
        "Tool inputs must match the specificity of the user's request. Ambiguous inputs "
        "must be flagged rather than silently assumed.",
        "Scope drift — dropping a state, country, or qualifier when calling a tool — "
        "produces results for the wrong target without any visible error signal.",
    ),
    (
        "REQ-09\nConfidence Calibration",
        "The agent's expressed certainty must be proportional to the strength of its tool "
        "evidence. Definitive language against empty or mismatched results must be flagged.",
        "False certainty causes users to take irreversible actions. Calibration errors are "
        "distinct from factual errors and require a separate check.",
    ),
    (
        "REQ-10\nSource Traceability",
        "Every specific entity must be traceable to a source span in the trace, not just "
        "plausible given the context.",
        "Provides full audit capability. Required for compliance, regulated industries, "
        "or post-incident investigation where provenance must be proven.",
    ),
]

for idx, (req, catches, why) in enumerate(t3_rows):
    row = t3_table.add_row()
    bg = COLOR_ROW1 if idx % 2 == 0 else COLOR_WHITE
    for ci, text in enumerate([req, catches, why]):
        cell = row.cells[ci]
        set_cell_bg(cell, bg)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR_BODY
        if ci == 0:
            run.bold = True
            run.font.color.rgb = COLOR_TEAL

for row in t3_table.rows:
    row.cells[0].width = Inches(1.4)
    row.cells[1].width = Inches(2.8)
    row.cells[2].width = Inches(2.2)

doc.add_paragraph()

# ── Trade-off note ─────────────────────────────────────────────────────────────
heading2("Trade-off and Risk")
body(
    "Running only Tier 1 will catch the most critical known failures — fabricated actions "
    "and misrepresented tool results. The residual risk is narrow but real: an agent can "
    "pass both gate checks and still introduce a scope error (REQ-08), a confidence "
    "overclaim (REQ-09), or a cross-turn contradiction (REQ-05) that causes harm. "
    "The decision of which tiers to run on which traffic should be proportional to the "
    "reversibility of the actions the agent is taking and the stakes of an incorrect outcome "
    "for the user."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
rb = p.add_run("Recommended minimum for production systems: ")
rb.bold = True
rb.font.size = Pt(10.5)
rb.font.color.rgb = COLOR_NAVY
rd = p.add_run(
    "Tier 1 on every trace. Tier 2 on every trace that passes Tier 1. "
    "Tier 3 on any trace where Tier 2 surfaces a flag, on all high-value transactions, "
    "and as a full-suite periodic audit on a sampled subset of traffic."
)
rd.font.size = Pt(10.5)
rd.font.color.rgb = COLOR_BODY

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# INDIVIDUAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════

heading1("Individual Requirements")

for req in requirements:
    req_id_para(req["id"], req["title"])

    field_label("Requirement")
    body(req["requirement"])

    field_label("Rationale")
    body(req["rationale"])

    field_label("Acceptance Criteria")
    for criterion in req["acceptance_criteria"]:
        bullet(criterion)

    field_label("Examples")
    add_examples_table(req["examples"])

    doc.add_paragraph()


# ── Footer note ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
add_paragraph_border_bottom(p, "C5CEDE")
p = doc.add_paragraph()
r = p.add_run(
    "Generated from analysis of Monocle/OpenTelemetry traces for the LG Travel Agent "
    "(test runs April 9–10, 2026).  All six traces were originally labelled no_hallucination; "
    "hallucinations were found in all six upon trace-level inspection."
)
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x7A, 0x8A, 0xA0)
r.italic = True

output_path = "/Users/careyjames/Documents/GitHub/lg-travel-agent/Hallucination_Evaluation_Requirements.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
