"""
Generate okahu_eval_test_scenarios.docx from okahu_eval_test_scenarios.json
"""
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AGENT_COLORS = {
    "lg_travel_agent":         RGBColor(0x1F, 0x49, 0x7D),  # dark blue
    "customer_care_agent":     RGBColor(0x37, 0x5C, 0x23),  # dark green
    "financial_services_agent": RGBColor(0x7B, 0x2C, 0x2C), # dark red
}

LABEL_COLORS = {
    # positive / clean outcomes
    "high_precision": RGBColor(0x2E, 0x86, 0x48),
    "unbiased": RGBColor(0x2E, 0x86, 0x48),
    "highly_relevant": RGBColor(0x2E, 0x86, 0x48),
    "no_misuse": RGBColor(0x2E, 0x86, 0x48),
    "completed": RGBColor(0x2E, 0x86, 0x48),
    "no_hallucination": RGBColor(0x2E, 0x86, 0x48),
    "correct": RGBColor(0x2E, 0x86, 0x48),
    "excellent_adherence": RGBColor(0x2E, 0x86, 0x48),
    "excellent": RGBColor(0x2E, 0x86, 0x48),
    "complete": RGBColor(0x2E, 0x86, 0x48),
    "excellent_retention": RGBColor(0x2E, 0x86, 0x48),
    "ok": RGBColor(0x2E, 0x86, 0x48),
    "no_pii": RGBColor(0x2E, 0x86, 0x48),
    "non_toxic": RGBColor(0x2E, 0x86, 0x48),
    "high_recall": RGBColor(0x2E, 0x86, 0x48),
    "yes": RGBColor(0x2E, 0x86, 0x48),
    "positive": RGBColor(0x2E, 0x86, 0x48),
    # warning / partial outcomes
    "medium_precision": RGBColor(0xB8, 0x86, 0x00),
    "potentially_biased": RGBColor(0xB8, 0x86, 0x00),
    "moderately_relevant": RGBColor(0xB8, 0x86, 0x00),
    "potential_misuse": RGBColor(0xB8, 0x86, 0x00),
    "partially_completed": RGBColor(0xB8, 0x86, 0x00),
    "minor_hallucination": RGBColor(0xB8, 0x86, 0x00),
    "partially_correct": RGBColor(0xB8, 0x86, 0x00),
    "good_adherence": RGBColor(0xB8, 0x86, 0x00),
    "good": RGBColor(0xB8, 0x86, 0x00),
    "mostly_complete": RGBColor(0xB8, 0x86, 0x00),
    "good_retention": RGBColor(0xB8, 0x86, 0x00),
    "potential_pii": RGBColor(0xB8, 0x86, 0x00),
    "mildly_toxic": RGBColor(0xB8, 0x86, 0x00),
    "medium_recall": RGBColor(0xB8, 0x86, 0x00),
    "idk": RGBColor(0xB8, 0x86, 0x00),
    "neutral": RGBColor(0xB8, 0x86, 0x00),
    "slightly_relevant": RGBColor(0xB8, 0x86, 0x00),
    "fair": RGBColor(0xB8, 0x86, 0x00),
    "partially_complete": RGBColor(0xB8, 0x86, 0x00),
    "not_attempted": RGBColor(0xB8, 0x86, 0x00),
    # negative / failure outcomes
    "low_precision": RGBColor(0xC0, 0x00, 0x00),
    "biased": RGBColor(0xC0, 0x00, 0x00),
    "irrelevant": RGBColor(0xC0, 0x00, 0x00),
    "clear_misuse": RGBColor(0xC0, 0x00, 0x00),
    "failed": RGBColor(0xC0, 0x00, 0x00),
    "major_hallucination": RGBColor(0xC0, 0x00, 0x00),
    "incorrect": RGBColor(0xC0, 0x00, 0x00),
    "poor_adherence": RGBColor(0xC0, 0x00, 0x00),
    "no_adherence": RGBColor(0xC0, 0x00, 0x00),
    "poor": RGBColor(0xC0, 0x00, 0x00),
    "incomplete": RGBColor(0xC0, 0x00, 0x00),
    "poor_retention": RGBColor(0xC0, 0x00, 0x00),
    "no_retention": RGBColor(0xC0, 0x00, 0x00),
    "frustrated": RGBColor(0xC0, 0x00, 0x00),
    "pii_leakage": RGBColor(0xC0, 0x00, 0x00),
    "moderately_toxic": RGBColor(0xC0, 0x00, 0x00),
    "highly_toxic": RGBColor(0xC0, 0x00, 0x00),
    "low_recall": RGBColor(0xC0, 0x00, 0x00),
    "no": RGBColor(0xC0, 0x00, 0x00),
    "negative": RGBColor(0xC0, 0x00, 0x00),
}

AGENT_DISPLAY = {
    "lg_travel_agent": "LG Travel Agent",
    "customer_care_agent": "Customer Care Agent",
    "financial_services_agent": "Financial Services Agent",
}

EVAL_ORDER = [
    "contextual_precision",
    "bias",
    "contextual_relevancy",
    "misuse",
    "mcp_task_completion",
    "hallucination",
    "argument_correctness",
    "role_adherence",
    "summarization",
    "conversation_completeness",
    "knowledge_retention",
    "frustration",
    "pii_leakage",
    "toxicity",
    "contextual_recall",
    "answer_relevancy",
    "sentiment",
]


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def rgb_to_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def add_colored_run(para, text, color=None, bold=False, size=None):
    run = para.add_run(text)
    if color:
        run.font.color.rgb = color
    if bold:
        run.font.bold = True
    if size:
        run.font.size = Pt(size)
    return run


def build_document(data):
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # --- Title ---
    title = doc.add_heading(data["document"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # --- Meta ---
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_colored_run(meta, f"Version {data['version']}  |  {data['date']}",
                    color=RGBColor(0x60, 0x60, 0x60), size=9)

    doc.add_paragraph(data["description"]).paragraph_format.space_after = Pt(4)

    # --- Agent legend ---
    doc.add_heading("Agents Under Test", level=1)
    for key, info in data["agents"].items():
        p = doc.add_paragraph(style="List Bullet")
        add_colored_run(p, AGENT_DISPLAY[key], color=AGENT_COLORS[key], bold=True)
        add_colored_run(p, f"  —  workflow: {info['workflow']}")

    # --- Color key ---
    doc.add_heading("Expected Outcome Color Key", level=1)
    key_para = doc.add_paragraph()
    for label, color in [
        ("Green = clean / passing outcome", RGBColor(0x2E, 0x86, 0x48)),
        ("  |  Amber = partial / warning outcome", RGBColor(0xB8, 0x86, 0x00)),
        ("  |  Red = failure / risk outcome", RGBColor(0xC0, 0x00, 0x00)),
    ]:
        add_colored_run(key_para, label, color=color, bold=True, size=9)

    doc.add_page_break()

    # --- Group scenarios by eval type ---
    by_eval: dict[str, list] = {e: [] for e in EVAL_ORDER}
    for s in data["scenarios"]:
        by_eval[s["eval_type"]].append(s)

    for eval_type in EVAL_ORDER:
        scenarios = by_eval.get(eval_type, [])
        if not scenarios:
            continue

        heading = doc.add_heading(eval_type.replace("_", " ").title(), level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        for s in scenarios:
            agent_key = s["agent"]
            agent_color = AGENT_COLORS.get(agent_key, RGBColor(0, 0, 0))
            outcome = s["expected_outcome"]
            outcome_color = LABEL_COLORS.get(outcome, RGBColor(0, 0, 0))

            # Scenario ID + agent header
            hdr = doc.add_paragraph()
            hdr.paragraph_format.space_before = Pt(6)
            hdr.paragraph_format.space_after = Pt(2)
            add_colored_run(hdr, f"{s['scenario_id']}", bold=True, size=11)
            add_colored_run(hdr, "  |  ")
            add_colored_run(hdr, AGENT_DISPLAY.get(agent_key, agent_key),
                            color=agent_color, bold=True, size=10)
            add_colored_run(hdr, f"  |  {s['session_type'].replace('_', ' ')}", size=9,
                            color=RGBColor(0x60, 0x60, 0x60))

            # Table: User Input | Eval | Expected Outcome | Tester Notes
            tbl = doc.add_table(rows=2, cols=4)
            tbl.style = "Table Grid"
            tbl.autofit = False

            # Column widths (total ~6.3 in usable)
            widths = [Inches(2.1), Inches(1.3), Inches(1.2), Inches(1.7)]
            for i, w in enumerate(widths):
                for row in tbl.rows:
                    row.cells[i].width = w

            # Header row
            headers = ["User Input", "Evaluation", "Expected Outcome", "Tester Notes"]
            for i, h in enumerate(headers):
                cell = tbl.rows[0].cells[i]
                set_cell_bg(cell, "1F497D")
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)

            # Data row
            data_row = tbl.rows[1]

            # User input — format turns
            ui_cell = data_row.cells[0]
            ui_cell.paragraphs[0].clear()
            turns = s["user_input"]
            for i, turn in enumerate(turns):
                p = ui_cell.paragraphs[0] if i == 0 else ui_cell.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                if len(turns) > 1:
                    add_colored_run(p, f"T{turn['turn']}: ",
                                    color=RGBColor(0x60, 0x60, 0x60), bold=True, size=8)
                add_colored_run(p, turn["message"], size=8)

            # Evaluation type
            eval_cell = data_row.cells[1]
            ep = eval_cell.paragraphs[0]
            add_colored_run(ep, eval_type, bold=True, size=8)

            # Expected outcome
            out_cell = data_row.cells[2]
            op = out_cell.paragraphs[0]
            add_colored_run(op, outcome, color=outcome_color, bold=True, size=9)

            # Tester notes
            note_cell = data_row.cells[3]
            np_ = note_cell.paragraphs[0]
            add_colored_run(np_, s["tester_notes"], size=8)

            doc.add_paragraph()  # spacer

        doc.add_page_break()

    return doc


if __name__ == "__main__":
    with open("okahu_eval_test_scenarios.json", "r") as f:
        data = json.load(f)

    doc = build_document(data)
    out = "okahu_eval_test_scenarios.docx"
    doc.save(out)
    print(f"Saved: {out}")
